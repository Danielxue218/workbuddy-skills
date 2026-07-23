#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""律师版研判报告 · Word(.docx) 生成器.

对标缪奇川合同审查技能的 Word 版式（封面大标题 + 报告信息表 + 分章 + 表格 + 品牌页脚），
**封面主色 = #991B1B 标准深红**（本 skill 唯一强调红）；功能色守本 skill 纪律：
决定性/最坏 = #991B1B、高前景 = #00A650、中前景 = #FF9900；**禁用 #FF0000**。
字体：中文宋体 + 西文 Times New Roman。

本文件是**参考实现**（内容为 demo-171），运行时按同一版式替换 REPORT 内容即可。
用法：python build_report_docx.py <输出.docx> [图目录(含 elements_matrix.png 等)]
"""
import sys, os
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RED   = RGBColor(0x99, 0x1B, 0x1B)   # 标准深红（封面 + 决定性/最坏）
INK   = RGBColor(0x1A, 0x1A, 0x1A)
INK2  = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x00, 0xA6, 0x50)   # 高前景
ORANGE= RGBColor(0xFF, 0x99, 0x00)   # 中前景
LINE  = "BFBFBF"
HEADFILL = "F1F5F9"
CN = "宋体"; EN = "Times New Roman"


def _set_cell_bg(cell, hexfill):
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:color'), 'auto'); sh.set(qn('w:fill'), hexfill)
    cell._tc.get_or_add_tcPr().append(sh)

def _para_shade(p, hexfill):
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:color'), 'auto'); sh.set(qn('w:fill'), hexfill)
    p._p.get_or_add_pPr().append(sh)

def _para_border(p, edges=('bottom',), sz=6, color=LINE):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    for e in edges:
        el = OxmlElement(f'w:{e}')
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '4'); el.set(qn('w:color'), color)
        pbdr.append(el)
    pPr.append(pbdr)

def _run(p, text, size=10.5, bold=False, color=INK, font=CN):
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.font.color.rgb = color
    r.font.name = font
    r._element.rPr.rFonts.set(qn('w:eastAsia'), CN)
    return r

def _table_borders(tbl):
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'4')
        el.set(qn('w:space'),'0'); el.set(qn('w:color'),LINE)
        borders.append(el)
    tblPr.append(borders)


def setup_styles(doc):
    st = doc.styles['Normal']
    st.font.name = EN; st.font.size = Pt(10.5); st.font.color.rgb = INK
    st.element.rPr.rFonts.set(qn('w:eastAsia'), CN)
    sec = doc.sections[0]
    sec.page_width = Mm(210); sec.page_height = Mm(297)
    sec.top_margin = Mm(22); sec.bottom_margin = Mm(20)
    sec.left_margin = Mm(24); sec.right_margin = Mm(24)


def add_footer(doc):
    p = doc.sections[0].footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "LEGAL AI TOOLMAKER · 法律工具制造者  |  缪奇川律师 出品", 8, False, INK2)


def cover(doc, meta):
    for _ in range(2): doc.add_paragraph()
    # 英文小标题
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "CLAIM-BASIS ANALYSIS REPORT", 12, True, RED, EN)
    # 深红横幅：中文主标题
    band = doc.add_paragraph(); band.alignment = WD_ALIGN_PARAGRAPH.CENTER
    band.paragraph_format.space_before = Pt(6); band.paragraph_format.space_after = Pt(6)
    _para_shade(band, "991B1B")
    r = _run(band, "请求权基础 · 要件审判研判报告", 22, True, WHITE)
    r2 = band.add_run(); 
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(sub, "（律师版 · 判决前冷推演）", 11, False, INK2)
    _para_border(sub, edges=('bottom',), sz=12, color="991B1B")
    for _ in range(2): doc.add_paragraph()

    # 报告信息表
    rows = [
        ("案　由", meta["案由"]),
        ("研判编号", meta["编号"]),
        ("研判立场", meta["立场"]),
        ("研判时态", meta["时态"]),
        ("经办律师", meta["律师"]),
        ("研判日期", meta["日期"]),
    ]
    tbl = doc.add_table(rows=len(rows), cols=2); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _table_borders(tbl)
    tbl.columns[0].width = Cm(4.5); tbl.columns[1].width = Cm(10.5)
    for i,(k,v) in enumerate(rows):
        c0, c1 = tbl.rows[i].cells
        c0.width = Cm(4.5); c1.width = Cm(10.5)
        _set_cell_bg(c0, HEADFILL)
        p0 = c0.paragraphs[0]; p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p0, k, 10.5, True, INK)
        p1 = c1.paragraphs[0]; _run(p1, v, 10.5, True, INK)
    for _ in range(3): doc.add_paragraph()
    foot = doc.add_paragraph(); foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_border(foot, edges=('top',), sz=6)
    _run(foot, "场景极度垂直 · SOP 极度精简 · 交付极度优雅", 9, False, INK2)
    doc.add_page_break()


def h1(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
    _para_border(p, edges=('bottom',), sz=8, color="991B1B")
    _run(p, text, 15, True, RED)

def h2(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    _run(p, text, 12, True, INK)

def para(doc, text, size=10.5):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
    _run(p, text, size, False, INK)
    return p

def bullet(doc, text, mark=None, mark_color=None):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.6); p.paragraph_format.space_after = Pt(2)
    if mark: _run(p, mark + " ", 10.5, True, mark_color or INK)
    _run(p, text, 10.5, False, INK)
    return p


def kv_table(doc, headers, rows, widths, decisive_col=None):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers)); _table_borders(tbl)
    for j,h in enumerate(headers):
        c = tbl.rows[0].cells[j]; _set_cell_bg(c, "991B1B")
        c.width = Cm(widths[j]); pp = c.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(pp, h, 10, True, WHITE)
    for i,row in enumerate(rows):
        decisive = decisive_col is not None and str(row[decisive_col]).strip().startswith("★")
        for j,val in enumerate(row):
            c = tbl.rows[1+i].cells[j]; c.width = Cm(widths[j])
            pp = c.paragraphs[0]
            col = RED if decisive else INK
            _run(pp, str(val), 9.5, decisive, col)
    return tbl


def build(out_path, img_dir):
    doc = Document(); setup_styles(doc); add_footer(doc)

    meta = {
        "案由": "建设工程施工合同纠纷（含建设工程价款优先受偿权）",
        "编号": "MQC-CB-20260708-171（demo）",
        "立场": "原告向（中天建设 · 承包人）",
        "时态": "前瞻（判决前冷推演）",
        "律师": "缪奇川 · 北京市东友律师事务所",
        "日期": "2026 年 7 月 8 日",
    }
    cover(doc, meta)

    # ── 一、研判基本信息 ──
    h1(doc, "一、研判基本信息")
    kv_table(doc, ["项目", "内容"], [
        ["当事人", "中天建设集团有限公司（承包人/原告） 诉 河南恒和置业有限公司（发包人/被告，破产中）"],
        ["请求权基础", "合同法286条→优先受偿权；合同无效但验收合格参照有效（民法典793）〔案时口径〕"],
        ["材料成熟度", "L2–L3（较完整卷宗）"],
        ["三条铁律", "分级不拒跑 · 三标出厂 · 客观与立场物理分离"],
    ], [3.5, 11.5])

    # ── 二、一页纸摘要 ──
    h1(doc, "二、一页纸摘要（结论先行）")
    p = para(doc, ""); _run(p, "【结论】", 11, True, RED)
    _run(p, "本案最可能"), _run(p, "部分支持", 10.5, True, INK)
    _run(p, "：工程款本金、利息、优先受偿权大概率获支持；停工损失恐被驳回。")
    h2(doc, "2.1  结果预判区间")
    bullet(doc, "判决支持 ／ 部分支持〔最可能〕 ／ 判决驳回诉请〔最坏〕 ／ 裁定驳回起诉〔概率低〕")
    bullet(doc, "最坏情形：结算不被采信→鉴定→数额落空，且优先受偿权除斥期间成立（叠加）。〔置信度：中〕", "■", RED)
    h2(doc, "2.2  胜负手清单")
    bullet(doc, "优先受偿权是否超 6 个月除斥期间（唯一决定性抗辩 · 全案红标）〔原告向｜置信度中〕", "★", RED)
    bullet(doc, "《审核报告》能否作结算依据（T4 数额）——本应第一决定性要件，经结算协议+印章自认前景转高", "■", ORANGE)
    bullet(doc, "破产债权确认之诉的程序前置抗辩（弱）", "■", GREEN)
    h2(doc, "2.3  核心建议")
    bullet(doc, "锁死结算协议成立（王峰波签字盖章+同一枚印章自认）→ 堵死对方造价鉴定申请〔解释二12条〕。")
    bullet(doc, "坐实优先受偿权在期限内主张（三次发函 + 执行程序视为行使，指导案例171号）。")

    # ── 三、六阶段研判（客观）──
    h1(doc, "三、六阶段研判（阶段一~五 · 客观，禁立场）")
    for t, body in [
        ("阶段一 · 接案与案情固定", "承包人/原告→原告向；对方发包人、破产中；无一审判决→前瞻；案由=建设工程施工合同纠纷；L2–L3。诉请均为给付之诉（界内），解除合同（形成之诉）已剔除。"),
        ("阶段二 · 找法", "请求权基础：合同法109·286；优先受偿权=合同法286+建工解释二21·22〔案时〕。〔版本时效〕民法典时代改民法典807+建工解释一35–42，期限6个月→18个月（41条）。抗辩二分：结算否认+鉴定/质量/程序前置（依职权审）· 除斥期间（须主张）。"),
        ("阶段三 · 证明前景（见附件图一）", "两层审查后进第二层。决定性抗辩=K1优先受偿权除斥期间（红标）。T4数额本应第一决定性要件，经修正因素（结算协议+印章自认）前景由存疑拉高至高，退出决定性之列。对方伪造公章抗辩=排除合理怀疑、前景低。"),
        ("阶段四 · 对抗与结果（见附件图二、图三）", "对手动作按危险度：否认审核报告+鉴定/除斥期间/质量/程序前置。裁判倾向：结算协议成立→不准鉴定（解释二12）；执行程序主张视为行使（171号）。结果区间见图三。"),
        ("阶段五 · 办案与检索（行动清单 · 仍客观）", "取证：结算对审记录/印章同一性/三次发函/施工月报。应对鉴定：以结算协议已成立抗辩。检索：解释二12条、指导案例171号。"),
    ]:
        h2(doc, t); para(doc, body)

    # 硬分隔
    sep = doc.add_paragraph(); sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep.paragraph_format.space_before = Pt(8); sep.paragraph_format.space_after = Pt(8)
    _para_shade(sep, "991B1B")
    _run(sep, "════  以下为立场打法，与上述客观研判物理分隔  ════", 10, True, WHITE)

    # ── 四、立场打法 ──
    h1(doc, "四、立场打法（阶段六 · 原告向）")
    bullet(doc, "主打：工程款本金+优先受偿权（以结算协议锁定数额）；备位：合同无效→参照有效结算。")
    bullet(doc, "攻防着力：锁死结算协议成立堵死鉴定（解释二12条）；优先受偿权以最有利起算点+行使方式主张（171号）。")
    bullet(doc, "时效/失权强制提醒：优先受偿权6个月除斥期间是我方最大失权风险，三次发函+执行程序主张两路并坐实。", "■", RED)
    bullet(doc, "成本收益（定性）：本金一环较稳、争点集中于优先受偿权，诉讼价值高，不宜轻易调解让步优先受偿权。")

    # ── 五、附件（三图 + 待核验）──
    doc.add_page_break()
    h1(doc, "五、附件")
    figs = [
        ("图一 · 要件-证据对照矩阵（④⑦）", "elements_matrix.png"),
        ("图二 · 争议焦点汇聚（⑥）", "issues_focus.png"),
        ("图三 · 结果预判区间带（⑨）", "outcome_band.png"),
    ]
    for title, fn in figs:
        h2(doc, title)
        fp = os.path.join(img_dir, fn) if img_dir else None
        if fp and os.path.exists(fp):
            pic = doc.add_paragraph(); pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pic.add_run().add_picture(fp, width=Cm(16))
        else:
            para(doc, f"（图未找到：{fn}）")
    h2(doc, "待核验清单（案时口径）")
    para(doc, "合同法109·286 ／ 建工解释（二）12·13·21·22 ／ 优先受偿权批复第4条 ／ 指导案例171号。现行对应（民法典793·807 ／ 建工解释一29·35·36·38·39·40·41）与 ✓已核/◐待核 见《条号核验》。输出前以核验为准。")

    # 免责
    dis = doc.add_paragraph(); dis.paragraph_format.space_before = Pt(10)
    _para_border(dis, edges=('top',), sz=6)
    _run(dis, "声明：本报告为代理立场下的判决前研判，非中立裁断、非判决模拟；结果涉法官自由裁量，不构成胜诉承诺。凡涉条号/案号标【待核验】者须核验后使用。", 9, False, INK2)

    doc.save(out_path)
    print("saved", out_path)


if __name__ == "__main__":
    out = sys.argv[1] if len(sys.argv) > 1 else "研判报告.docx"
    imgs = sys.argv[2] if len(sys.argv) > 2 else ""
    build(out, imgs)
