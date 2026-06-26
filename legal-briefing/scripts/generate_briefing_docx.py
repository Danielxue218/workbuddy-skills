#!/usr/bin/env python3
"""Generate a professional Word document (.docx) from the legal briefing JSON.

Fonts: 宋体 (SimSun) + Times New Roman for ASCII.
Base size: 小四 (12pt).
Paragraph spacing: 段前0.5行, 段后0.5行.
行距: 1.15 倍行距."""

import json
import os
import sys

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ------ Config ------
# Support command-line arguments:
# sys.argv[1] = JSON path (required)
# sys.argv[2] = OUTPUT_DIR (optional, defaults to same dir as JSON)
if len(sys.argv) > 1:
    JSON_PATH = sys.argv[1]
    if len(sys.argv) > 2:
        OUTPUT_DIR = sys.argv[2]
    else:
        OUTPUT_DIR = os.path.dirname(JSON_PATH)
else:
    JSON_PATH = r"E:\workbuddy\briefings\2026-06-20-briefing-v3.json"
    OUTPUT_DIR = r"E:\workbuddy\briefings"

FONT_CN = '宋体'
FONT_EN = 'Times New Roman'

# Chinese font sizes: 小二=18pt, 四号=14pt, 小四=12pt, 五号=10.5pt
SIZE_TITLE = Pt(18)         # 小二 - 主标题
SIZE_BODY = Pt(12)          # 小四 - 正文
SIZE_META = Pt(10.5)        # 五号 - 摘要信息
SIZE_ART_TITLE = Pt(14)     # 四号 - 文章标题
SIZE_FOOTER = Pt(10.5)      # 五号 - 页脚

# 段前段后 0.5 行（小四 12pt × 0.5 = 6pt）
SPACING = Pt(6)

# Colors
PRUSSIAN_BLUE = RGBColor(0x1A, 0x25, 0x30)
MUTED_GOLD = RGBColor(0xC5, 0xA0, 0x59)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
LIGHT_GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def set_font(run, size=SIZE_BODY, bold=False, color=DARK_GRAY):
    """Set 宋体 + Times New Roman for a run."""
    run.font.size = size
    run.bold = bold
    run.font.color.rgb = color
    run.font.name = FONT_EN
    rPr = run._element.get_or_add_rPr()
    rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rf)
    rf.set(qn('w:eastAsia'), FONT_CN)
    rf.set(qn('w:ascii'), FONT_EN)
    rf.set(qn('w:hAnsi'), FONT_EN)


def set_spacing(para, space_before=SPACING, space_after=SPACING):
    """段前0.5行, 段后0.5行, 1.15倍行距."""
    para.paragraph_format.space_before = space_before
    para.paragraph_format.space_after = space_after
    para.paragraph_format.line_spacing = 1.15


def set_shading(para, color_hex):
    """深蓝底."""
    pPr = para._element.get_or_add_pPr()
    s = pPr.makeelement(
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd',
        {
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val': 'clear',
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color': 'auto',
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill': color_hex,
        },
    )
    pPr.append(s)


def add_p(doc, text, size=SIZE_BODY, bold=False, color=DARK_GRAY,
          alignment=None, shading=None):
    """Add styled paragraph."""
    p = doc.add_paragraph()
    set_spacing(p)
    if alignment is not None:
        p.alignment = alignment
    if shading:
        set_shading(p, shading)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return p


def add_hr(doc):
    """Gold horizontal rule."""
    p = doc.add_paragraph()
    set_spacing(p, Pt(3), Pt(3))
    pPr = p._element.get_or_add_pPr()
    pb = pPr.makeelement(
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pBdr', {}
    )
    bt = pb.makeelement(
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bottom',
        {
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val': 'single',
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz': '6',
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}space': '1',
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color': 'C5A059',
        },
    )
    pb.append(bt)
    pPr.append(pb)


def generate_docx():
    with open(JSON_PATH, 'r', encoding='utf-8') as f:
        data = json.load(f)

    doc = Document()

    # Default style
    style = doc.styles['Normal']
    style.font.name = FONT_EN
    style.font.size = SIZE_BODY
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CN)

    # Page margins (Word 默认)
    for sec in doc.sections:
        sec.top_margin = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(3.17)
        sec.right_margin = Cm(3.17)

    # ===== Header =====
    add_p(doc, 'WorkBuddy 今日法律要闻简报',
          SIZE_TITLE, bold=True, color=PRUSSIAN_BLUE,
          alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, f"({data['date']})",
          SIZE_BODY, color=LIGHT_GRAY,
          alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_hr(doc)

    # ===== Summary =====
    info = (f"共检索 {data['total_fetched']} 篇 | "
            f"法务筛选 {data['after_filter']} 篇 | "
            f"质量优先入选 {data['total_selected']} 篇")
    add_p(doc, info, SIZE_META, color=LIGHT_GRAY,
          alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, f"选取原则：{data['selection_principle']}",
          SIZE_META, color=MUTED_GOLD,
          alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, f"方向覆盖：{' · '.join(data['directions_covered'])}",
          SIZE_META, color=LIGHT_GRAY,
          alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_hr(doc)

    # ===== No-coverage =====
    add_p(doc, '以下方向今日无高价值内容入选（按质量优先原则，不予降级）：',
          SIZE_META, color=LIGHT_GRAY)
    for dn, rr in data['directions_no_coverage'].items():
        add_p(doc, f'· {dn}：{rr}', SIZE_META, color=LIGHT_GRAY)
    add_hr(doc)

    # ===== Articles =====
    for art in data['selected']:
        # Title bar (blue bg, white text)
        pt = doc.add_paragraph()
        set_spacing(pt, Pt(12))
        set_shading(pt, '1A2530')
        r = pt.add_run(f"【{art['index']}】{art['title']}")
        r.font.size = SIZE_ART_TITLE
        r.bold = True
        r.font.color.rgb = WHITE
        r.font.name = FONT_EN
        rPr = r._element.get_or_add_rPr()
        rf = rPr.find(qn('w:rFonts'))
        if rf is None:
            rf = rPr.makeelement(qn('w:rFonts'), {})
            rPr.insert(0, rf)
        rf.set(qn('w:eastAsia'), FONT_CN)
        rf.set(qn('w:ascii'), FONT_EN)
        rf.set(qn('w:hAnsi'), FONT_EN)

        # Meta
        parts = [f"来源：{art['source']}"]
        if art.get('date'):
            parts.append(f"日期：{art['date']}")
        if art.get('case_number'):
            parts.append(f"案号：{art['case_number']}")
        if art.get('directions'):
            parts.append(f"方向：{' · '.join(art['directions'])}")
        add_p(doc, '  |  '.join(parts), SIZE_META, color=LIGHT_GRAY)

        # Focus
        pf = doc.add_paragraph()
        set_spacing(pf)
        rl = pf.add_run('争议焦点：')
        set_font(rl, SIZE_BODY, bold=True, color=MUTED_GOLD)
        rc = pf.add_run(art['focus'])
        set_font(rc, SIZE_BODY, color=DARK_GRAY)

        # Ruling
        pr = doc.add_paragraph()
        set_spacing(pr)
        rl = pr.add_run('法院/仲裁裁判：')
        set_font(rl, SIZE_BODY, bold=True, color=MUTED_GOLD)
        rc = pr.add_run(art['ruling'])
        set_font(rc, SIZE_BODY, color=DARK_GRAY)

        # Opportunity
        po = doc.add_paragraph()
        set_spacing(po)
        rl = po.add_run('应用场景与案源价值：')
        set_font(rl, SIZE_BODY, bold=True, color=MUTED_GOLD)
        rc = po.add_run(art['opportunity'])
        set_font(rc, SIZE_BODY, color=DARK_GRAY)

    # ===== Footer =====
    add_hr(doc)
    add_p(doc, '薛龙 | 合伙人/律师 | 18016302187',
          SIZE_FOOTER, color=LIGHT_GRAY,
          alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, '—— 由 WorkBuddy AI 自动生成 ——',
          SIZE_FOOTER, color=LIGHT_GRAY,
          alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Save
    output_path = os.path.join(OUTPUT_DIR, f"{data['date']}-法律简报.docx")
    doc.save(output_path)
    print(f"[OK] {output_path}")
    return output_path


if __name__ == '__main__':
    generate_docx()
