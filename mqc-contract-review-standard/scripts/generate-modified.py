#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
generate-modified.py · v1.0.0(随 skill v2.0.0 首发)

四模式合同修改文件生成器 —— 在原合同 docx 上直接施工,生成:

    模式        输出文件              说明
    --------    ------------------    ----------------------------------------
    annotate    {合同名}_批注版.docx      不改正文,逐条添加 Word 批注(审查意见)
    revise      {合同名}_修订版.docx      修订模式(Track Changes)修改,无批注
    both        {合同名}_修订批注版.docx   修订模式修改 + 批注,二合一
    clean       {合同名}_清洁版.docx      直接改完,无修改痕迹,可直接签订
    all         一次生成以上四件

数据来源:复用 report-data.json(与 generate-report.py 同一份数据),
依赖其中每条 finding 的三个字段:
    original_text        批注锚点(合同原文逐字引用)
    suggestion_original   修订/清洁版中被替换的原文(必须与合同逐字一致)
    suggestion_revised    替换后的文字(`**...**` 标记在本脚本中会被剥离)
以及 missing_clauses 的 bridge_clause(修订/清洁版自动补入为新条款)。

署名规则:批注作者与修订作者 = data["lawyer"](经办律师),不使用工具作者名。

定位引擎(三层防御):
    1. 精确匹配 · 全半角标点归一化 + 空白剥离后精确查找
    2. 模糊匹配 · difflib 相似度兜底(阈值 0.80),命中后按匹配块还原原文区间
    3. 显式失败 · 两层都失败的条目列入"定位失败清单",绝不静默跳过

用法:
    # 一次生成四件(推荐,Step 4 默认路径)
    python3 scripts/generate-modified.py \\
        --contract "原合同.docx" \\
        --data "report-data.json" \\
        --mode all \\
        --output-dir "输出目录" \\
        --contract-name "技术服务合同"

    # 单独生成某一版
    python3 scripts/generate-modified.py \\
        --contract "原合同.docx" --data "report-data.json" \\
        --mode revise --output "技术服务合同_修订版.docx"

依赖:python-docx >= 1.2.0(批注 API);修订模式通过 OOXML w:ins/w:del 实现。

作者:缪奇川(Miao Qichuan) · mqc-contract-review-standard v2.0.0
"""

import argparse
import copy
import datetime
import difflib
import json
import re
import sys
from pathlib import Path

try:
    import docx  # noqa: F401
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.text.run import Run
except ImportError:
    print("错误:未安装 python-docx。请执行 pip install python-docx>=1.2.0", file=sys.stderr)
    sys.exit(1)


# ============================================================
# 常量
# ============================================================

MODE_FILENAMES = {
    "annotate": "批注版",
    "revise": "修订版",
    "both": "修订批注版",
    "clean": "清洁版",
}

FUZZY_THRESHOLD = 0.80          # 模糊匹配接受阈值
BOLD_MARK_RE = re.compile(r"\*\*(.+?)\*\*", re.S)

# 签章页/正文结束锚点(缺失条款插入在此之前)
SIGNATURE_ANCHOR_RE = re.compile(
    r"以下无正文|签章页|签字页|盖章页|^甲\s*方\s*[(（]\s*盖章|^甲\s*方\s*[:：].*盖章"
)

# 条款标题:第X条(中文数字)
CLAUSE_HEADING_RE = re.compile(r"^第([零一二三四五六七八九十百]+)条")

CN_DIGITS = {"零": 0, "一": 1, "二": 2, "三": 3, "四": 4,
             "五": 5, "六": 6, "七": 7, "八": 8, "九": 9}


# ============================================================
# 中文数字互转(支持 1-99,覆盖常规合同条款数)
# ============================================================

def cn_to_int(s):
    """'十五' -> 15;'二十三' -> 23;失败返回 None。"""
    s = s.strip()
    if not s:
        return None
    if s == "十":
        return 10
    if "十" in s:
        parts = s.split("十")
        tens = CN_DIGITS.get(parts[0], 1) if parts[0] else 1
        ones = CN_DIGITS.get(parts[1], 0) if len(parts) > 1 and parts[1] else 0
        return tens * 10 + ones
    total = 0
    for ch in s:
        if ch not in CN_DIGITS:
            return None
        total = total * 10 + CN_DIGITS[ch]
    return total


def int_to_cn(n):
    """16 -> '十六';23 -> '二十三'。"""
    units = "零一二三四五六七八九"
    if n < 10:
        return units[n]
    if n < 20:
        return "十" + (units[n % 10] if n % 10 else "")
    tens, ones = divmod(n, 10)
    return units[tens] + "十" + (units[ones] if ones else "")


# ============================================================
# 文本归一化(匹配用,不影响文档内容)
# ============================================================

_PUNCT_PAIRS = {
    "，": ",", "。": ".", "：": ":", "；": ";", "！": "!", "？": "?",
    "（": "(", "）": ")", "【": "[", "】": "]", "「": '"', "」": '"',
    "\u201c": '"', "\u201d": '"', "\u2018": "'", "\u2019": "'",
    "、": ",", "－": "-", "—": "-", "～": "~", "％": "%", "．": ".",
}


def normalize_with_map(text):
    """归一化文本用于匹配。

    返回 (norm_str, idx_map):norm_str 的第 i 个字符对应原文 text[idx_map[i]]。
    规则:剥离所有空白字符;全角标点折半角;其余字符原样。
    """
    chars, idx_map = [], []
    for i, ch in enumerate(text):
        if ch.isspace() or ch in ("\u200b", "\ufeff"):
            continue
        chars.append(_PUNCT_PAIRS.get(ch, ch))
        idx_map.append(i)
    return "".join(chars), idx_map


def strip_bold_marks(text):
    """剥离 **...** 标记,返回纯文本(修订/清洁版写入合同的文字)。"""
    return BOLD_MARK_RE.sub(r"\1", text or "")


# ============================================================
# 定位引擎
# ============================================================

def para_runs_text(para):
    """段落内按 run 拼接的文本(与 run 偏移一致的坐标系)。"""
    return "".join(r.text or "" for r in para.runs)


def locate_in_paragraph(para, target):
    """在段落中定位 target。

    返回 (raw_start, raw_end, method) 或 None。
    raw_* 是相对 para_runs_text(para) 的偏移;method ∈ {'exact', 'fuzzy'}。
    """
    raw = para_runs_text(para)
    if not raw.strip():
        return None
    norm_raw, idx_map = normalize_with_map(raw)
    norm_tgt, _ = normalize_with_map(target)
    if not norm_tgt:
        return None

    # 第一层:归一化精确匹配
    pos = norm_raw.find(norm_tgt)
    if pos >= 0:
        raw_start = idx_map[pos]
        raw_end = idx_map[pos + len(norm_tgt) - 1] + 1
        return raw_start, raw_end, "exact"

    # 第二层:模糊匹配(要求段落与目标有足够重叠)
    sm = difflib.SequenceMatcher(None, norm_raw, norm_tgt, autojunk=False)
    blocks = [b for b in sm.get_matching_blocks() if b.size > 0]
    if not blocks:
        return None
    matched = sum(b.size for b in blocks)
    ratio = matched / max(len(norm_tgt), 1)
    if ratio < FUZZY_THRESHOLD:
        return None
    span_start = min(b.a for b in blocks)
    last = max(blocks, key=lambda b: b.a + b.size)
    span_end = last.a + last.size
    raw_start = idx_map[span_start]
    raw_end = idx_map[span_end - 1] + 1
    return raw_start, raw_end, "fuzzy"


def locate_in_document(doc, target):
    """全文档定位,返回 (paragraph, raw_start, raw_end, method) 或 None。

    优先精确命中;全文无精确命中时,取模糊匹配得分最高的段落。
    """
    best = None
    best_ratio = 0.0
    norm_tgt, _ = normalize_with_map(target)
    for para in doc.paragraphs:
        res = locate_in_paragraph(para, target)
        if res is None:
            continue
        raw_start, raw_end, method = res
        if method == "exact":
            return para, raw_start, raw_end, "exact"
        # 模糊命中:记录得分最高者
        norm_seg, _m = normalize_with_map(
            para_runs_text(para)[raw_start:raw_end])
        ratio = difflib.SequenceMatcher(
            None, norm_seg, norm_tgt, autojunk=False).ratio()
        if ratio > best_ratio:
            best_ratio = ratio
            best = (para, raw_start, raw_end, "fuzzy")
    return best


# ============================================================
# run 切分与隔离
# ============================================================

def _set_run_text_preserve(r_element, text):
    """把 w:r 元素的文本内容替换为 text(保留 rPr)。"""
    for child in list(r_element):
        if child.tag != qn("w:rPr"):
            r_element.remove(child)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r_element.append(t)


def split_run_at(para, run, offset):
    """把 run 在 offset 处一分为二,返回 (前半 Run, 后半 Run)。"""
    text = run.text or ""
    r = run._r
    new_r = copy.deepcopy(r)
    r.addnext(new_r)
    _set_run_text_preserve(r, text[:offset])
    _set_run_text_preserve(new_r, text[offset:])
    return Run(r, para), Run(new_r, para)


def isolate_runs(para, raw_start, raw_end):
    """切分段落 run,使 [raw_start, raw_end) 恰好被若干完整 run 覆盖。

    返回覆盖该区间的 Run 列表(顺序)。
    """
    # 起点切分
    cursor = 0
    runs = list(para.runs)
    i = 0
    while i < len(runs):
        ln = len(runs[i].text or "")
        if cursor + ln > raw_start:
            if cursor < raw_start:
                _, right = split_run_at(para, runs[i], raw_start - cursor)
                runs = list(para.runs)
                cursor = raw_start
                # right 即区间起始 run,重新定位索引
                i = next(idx for idx, r in enumerate(runs)
                         if r._r is right._r)
            break
        cursor += ln
        i += 1
    # 终点切分
    target_runs = []
    runs = list(para.runs)
    cursor = 0
    for r in runs:
        ln = len(r.text or "")
        if cursor >= raw_end:
            break
        if cursor + ln <= raw_start:
            cursor += ln
            continue
        # r 与区间有交叠
        if cursor + ln > raw_end:
            left, _right = split_run_at(para, r, raw_end - cursor)
            target_runs.append(left)
            break
        target_runs.append(r)
        cursor += ln
    return target_runs


# ============================================================
# OOXML 修订(Track Changes)原语
# ============================================================

class RevisionIdAllocator:
    def __init__(self, start=1000):
        self._next = start

    def new(self):
        self._next += 1
        return str(self._next)


def _make_tracked_container(tag, rid, author, date_iso):
    el = OxmlElement(tag)
    el.set(qn("w:id"), rid)
    el.set(qn("w:author"), author)
    el.set(qn("w:date"), date_iso)
    return el


def wrap_runs_as_deletion(runs, rid_alloc, author, date_iso):
    """把若干连续 run 包裹进 <w:del>,w:t 转 w:delText。返回 del 元素。"""
    first_r = runs[0]._r
    del_el = _make_tracked_container("w:del", rid_alloc.new(), author, date_iso)
    first_r.addprevious(del_el)
    for run in runs:
        r = run._r
        del_el.append(r)  # 移动
        for t in r.findall(qn("w:t")):
            t.tag = qn("w:delText")
    return del_el


def make_insertion(text, rpr_template, rid_alloc, author, date_iso):
    """构造 <w:ins> 包裹的新 run。返回 (ins 元素, 新 w:r 元素)。"""
    ins_el = _make_tracked_container("w:ins", rid_alloc.new(), author, date_iso)
    new_r = OxmlElement("w:r")
    if rpr_template is not None:
        new_r.append(copy.deepcopy(rpr_template))
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    new_r.append(t)
    ins_el.append(new_r)
    return ins_el, new_r


def mark_paragraph_inserted(para, rid_alloc, author, date_iso):
    """把整段(含段落标记)标记为修订插入:run 包 w:ins + 段落标记 rPr/w:ins。"""
    p = para._p
    # 段落标记
    pPr = p.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p.insert(0, pPr)
    rPr = pPr.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        pPr.append(rPr)
    rPr.insert(0, _make_tracked_container(
        "w:ins", rid_alloc.new(), author, date_iso))
    # 全部 run 包 w:ins
    runs = list(para.runs)
    if runs:
        ins_el = _make_tracked_container(
            "w:ins", rid_alloc.new(), author, date_iso)
        runs[0]._r.addprevious(ins_el)
        for run in runs:
            ins_el.append(run._r)


# ============================================================
# 批注文案
# ============================================================

def build_comment_text(finding):
    """按视觉规范组装批注文案:■ 功能符号 + 文字标签,不用 emoji。"""
    risk = (finding.get("risk_level") or "").strip()
    tags = [t for t in [
        risk,
        finding.get("checklist_id"),
        finding.get("confirmation_level"),
        "精准审查" if finding.get("is_precise_review") else None,
    ] if t]
    header = "【" + " · ".join(tags) + "】" + (finding.get("title") or "")
    lines = [header]
    if finding.get("risk_description"):
        lines.append("风险说明:" + finding["risk_description"])
    if finding.get("suggestion_revised"):
        lines.append("修改建议:" + strip_bold_marks(finding["suggestion_revised"]))
    if finding.get("related_laws"):
        lines.append("相关法条:" + finding["related_laws"])
    return "\n".join(lines)


def build_extra_comment_text(extra):
    lines = ["【清单外发现】" + (extra.get("title") or "")]
    if extra.get("analysis"):
        lines.append("问题分析:" + extra["analysis"])
    if extra.get("suggestion"):
        lines.append("建议处理:" + strip_bold_marks(extra["suggestion"]))
    return "\n".join(lines)


def build_missing_comment_text(missing_list):
    lines = ["【缺失条款提示】本合同经对照审查清单,缺失以下应包含的条款:"]
    for i, m in enumerate(missing_list, 1):
        lines.append(f"{i}. {m.get('clause', '')}"
                     f"({m.get('checklist_id', '')} · {m.get('importance', '')})"
                     f" —— 建议条款:{strip_bold_marks(m.get('bridge_clause', ''))}")
    return "\n".join(lines)


def add_comment_on_runs(doc, runs, text, author, initials):
    """python-docx 1.2.0+ 批注 API 封装。"""
    return doc.add_comment(runs=runs, text=text,
                           author=author, initials=initials or "")


# ============================================================
# 缺失条款插入
# ============================================================

def find_insert_anchor(doc):
    """返回缺失条款应插入其前的段落(签章页标记段);找不到则返回 None。"""
    for para in doc.paragraphs:
        if SIGNATURE_ANCHOR_RE.search(para.text or ""):
            return para
    return None


def find_max_clause_number(doc):
    """扫描'第X条'标题,返回 (最大条号 int, 模板标题段, 模板正文段)。"""
    max_num, heading_para, last_heading_idx = 0, None, -1
    paras = doc.paragraphs
    for i, para in enumerate(paras):
        m = CLAUSE_HEADING_RE.match((para.text or "").strip())
        if m:
            num = cn_to_int(m.group(1))
            if num and num > max_num:
                max_num = num
                heading_para = para
                last_heading_idx = i
    body_para = None
    if last_heading_idx >= 0:
        for para in paras[last_heading_idx + 1:]:
            if (para.text or "").strip():
                body_para = para
                break
    return max_num, heading_para, body_para


def clone_paragraph_before(anchor_para, template_para, text):
    """在 anchor 前插入一个段落,克隆 template 的段落格式和 run 格式。"""
    new_p = copy.deepcopy(template_para._p)
    # 清空克隆段中除 pPr 外的内容
    for child in list(new_p):
        if child.tag != qn("w:pPr"):
            new_p.remove(child)
    # 用模板首个 run 的 rPr 构造新 run
    rpr_src = None
    for r in template_para._p.findall(qn("w:r")):
        rpr = r.find(qn("w:rPr"))
        if rpr is not None:
            rpr_src = rpr
        break
    new_r = OxmlElement("w:r")
    if rpr_src is not None:
        new_r.append(copy.deepcopy(rpr_src))
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    new_r.append(t)
    new_p.append(new_r)
    anchor_para._p.addprevious(new_p)
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, anchor_para._parent)


def insert_missing_clauses(doc, missing_list, as_revision,
                           rid_alloc, author, date_iso, report):
    """把缺失条款按'第X条'编号补入正文末尾(签章页之前)。"""
    if not missing_list:
        return
    anchor = find_insert_anchor(doc)
    max_num, heading_tpl, body_tpl = find_max_clause_number(doc)
    if anchor is None or heading_tpl is None:
        report["missing_failed"] = [m.get("clause") for m in missing_list]
        return
    body_tpl = body_tpl or heading_tpl
    for m in missing_list:
        max_num += 1
        clause_title = (m.get("clause") or "").replace("条款", "").strip()
        heading_text = f"第{int_to_cn(max_num)}条  {clause_title}"
        body_text = f"{max_num}.1  {strip_bold_marks(m.get('bridge_clause', ''))}"
        h_para = clone_paragraph_before(anchor, heading_tpl, heading_text)
        b_para = clone_paragraph_before(anchor, body_tpl, body_text)
        if as_revision:
            mark_paragraph_inserted(h_para, rid_alloc, author, date_iso)
            mark_paragraph_inserted(b_para, rid_alloc, author, date_iso)
        report["missing_inserted"].append(
            f"{m.get('clause')}(补入为 第{int_to_cn(max_num)}条)")


# ============================================================
# 主处理流程
# ============================================================

def parse_review_date(data):
    """把 '2026 年 4 月 17 日' 解析为修订元数据用的 ISO 时间。"""
    s = data.get("review_date") or ""
    m = re.search(r"(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日", s)
    if m:
        return "{:04d}-{:02d}-{:02d}T00:00:00Z".format(
            int(m.group(1)), int(m.group(2)), int(m.group(3)))
    return datetime.datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")


def derive_initials(name):
    return (name or "")[:2] or "审"


def modifiable_findings(data):
    """筛选可用于修订/清洁版的 finding:双字段齐备且未显式关闭。"""
    out = []
    for f in data.get("findings", []):
        if not f.get("apply_modification", True):
            continue
        if f.get("suggestion_original") and f.get("suggestion_revised"):
            out.append(f)
    return out


def apply_mode(contract_path, data, mode, output_path):
    """在原合同副本上执行一种模式,返回处理报告 dict。"""
    doc = Document(str(contract_path))
    lawyer = (data.get("lawyer") or "").strip() or "审查律师"
    initials = derive_initials(lawyer)
    date_iso = parse_review_date(data)
    rid_alloc = RevisionIdAllocator()

    do_comment = mode in ("annotate", "both")
    do_revise = mode in ("revise", "both")
    do_clean = mode == "clean"

    report = {
        "mode": mode, "located": [], "fuzzy": [], "failed": [],
        "missing_inserted": [], "missing_failed": [], "skipped": [],
    }

    # ---------- 逐条 finding ----------
    for f in data.get("findings", []):
        label = f"{f.get('checklist_id', '?')} {f.get('title', '')}"

        # 修订 / 清洁:替换 suggestion_original → suggestion_revised
        if do_revise or do_clean:
            if not (f.get("suggestion_original") and f.get("suggestion_revised")):
                report["skipped"].append(label + "(无替换文本对)")
            elif not f.get("apply_modification", True):
                report["skipped"].append(label + "(apply_modification=false)")
            else:
                loc = locate_in_document(doc, f["suggestion_original"])
                if loc is None:
                    report["failed"].append(label)
                else:
                    para, s, e, method = loc
                    runs = isolate_runs(para, s, e)
                    new_text = strip_bold_marks(f["suggestion_revised"])
                    if do_clean:
                        runs[0].text = new_text
                        for r in runs[1:]:
                            r._r.getparent().remove(r._r)
                    else:
                        rpr_tpl = runs[0]._r.find(qn("w:rPr"))
                        del_el = wrap_runs_as_deletion(
                            runs, rid_alloc, lawyer, date_iso)
                        ins_el, ins_r = make_insertion(
                            new_text, rpr_tpl, rid_alloc, lawyer, date_iso)
                        del_el.addnext(ins_el)
                        # both 模式:批注锚在新插入的文字上
                        if do_comment:
                            add_comment_on_runs(
                                doc, [Run(ins_r, para)],
                                build_comment_text(f), lawyer, initials)
                    (report["fuzzy"] if method == "fuzzy"
                     else report["located"]).append(label)
                    continue  # both 模式下批注已随修订完成

        # 纯批注(annotate 模式;both 模式定位失败时也退回锚原文批注)
        if do_comment:
            anchor_text = f.get("original_text") or f.get("suggestion_original")
            if not anchor_text:
                report["skipped"].append(label + "(无锚点原文)")
                continue
            loc = locate_in_document(doc, anchor_text)
            if loc is None:
                report["failed"].append(label)
                continue
            para, s, e, method = loc
            runs = isolate_runs(para, s, e)
            add_comment_on_runs(doc, runs, build_comment_text(f),
                                lawyer, initials)
            if mode == "annotate":
                (report["fuzzy"] if method == "fuzzy"
                 else report["located"]).append(label)

    # ---------- 清单外发现:仅批注 ----------
    if do_comment:
        for extra in data.get("extra_findings", []):
            label = "清单外 · " + (extra.get("title") or "")
            anchor_text = extra.get("original_text")
            if not anchor_text:
                report["skipped"].append(label + "(无锚点原文)")
                continue
            loc = locate_in_document(doc, anchor_text)
            if loc is None:
                report["failed"].append(label)
                continue
            para, s, e, _m = loc
            runs = isolate_runs(para, s, e)
            add_comment_on_runs(doc, runs, build_extra_comment_text(extra),
                                lawyer, initials)

    # ---------- 缺失条款 ----------
    missing = data.get("missing_clauses", [])
    if missing:
        if do_revise or do_clean:
            insert_missing_clauses(doc, missing, as_revision=do_revise,
                                   rid_alloc=rid_alloc, author=lawyer,
                                   date_iso=date_iso, report=report)
        elif mode == "annotate":
            # 批注版:在最后一个条款正文上锚一条"缺失条款提示"批注
            _n, _h, body_tpl = find_max_clause_number(doc)
            anchor_para = body_tpl
            if anchor_para is not None and anchor_para.runs:
                add_comment_on_runs(doc, list(anchor_para.runs)[-1:],
                                    build_missing_comment_text(missing),
                                    lawyer, initials)

    # ---------- 元数据与保存 ----------
    try:
        doc.core_properties.last_modified_by = lawyer
        doc.core_properties.comments = (
            f"mqc-contract-review-standard v2.0.0 · "
            f"{MODE_FILENAMES.get(mode, mode)} · 经办律师:{lawyer}")
    except Exception:
        pass
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    report["output"] = str(output_path)
    return report


def print_report(report):
    mode_cn = MODE_FILENAMES.get(report["mode"], report["mode"])
    print(f"\n[{mode_cn}] 已生成:{report['output']}")
    if report["located"]:
        print(f"  精确定位 {len(report['located'])} 处:"
              + ";".join(report["located"]))
    if report["fuzzy"]:
        print(f"  ⚠ 模糊定位 {len(report['fuzzy'])} 处(建议人工核对):"
              + ";".join(report["fuzzy"]))
    if report["missing_inserted"]:
        print("  缺失条款已补入:" + ";".join(report["missing_inserted"]))
    if report["missing_failed"]:
        print("  ⚠ 缺失条款插入失败(未找到条款结构/签章锚点):"
              + ";".join(report["missing_failed"]))
    if report["skipped"]:
        print("  跳过:" + ";".join(report["skipped"]))
    if report["failed"]:
        print(f"  ✗ 定位失败 {len(report['failed'])} 处(需人工处理):"
              + ";".join(report["failed"]))


def main():
    parser = argparse.ArgumentParser(
        description="四模式合同修改文件生成器(批注/修订/修订批注/清洁)")
    parser.add_argument("--contract", required=True, help="原合同 .docx 路径")
    parser.add_argument("--data", required=True,
                        help="report-data.json 路径(复用审查报告数据)")
    parser.add_argument("--mode", required=True,
                        choices=["annotate", "revise", "both", "clean", "all"])
    parser.add_argument("--output", help="输出 .docx 路径(单模式)")
    parser.add_argument("--output-dir", help="输出目录(--mode all)")
    parser.add_argument("--contract-name",
                        help="合同名(--mode all 时用于命名,缺省取 JSON 的 contract_name)")
    args = parser.parse_args()

    contract = Path(args.contract)
    if not contract.exists():
        print(f"错误:原合同文件不存在:{contract}", file=sys.stderr)
        sys.exit(1)
    if contract.suffix.lower() != ".docx":
        print("错误:--contract 必须是 .docx 文件(扫描版 PDF 请先转换)",
              file=sys.stderr)
        sys.exit(1)

    with open(args.data, encoding="utf-8") as fh:
        data = json.load(fh)

    if args.mode == "all":
        out_dir = Path(args.output_dir or ".")
        name = args.contract_name or data.get("contract_name") or contract.stem
        any_failed = False
        for mode, suffix in MODE_FILENAMES.items():
            out = out_dir / f"{name}_{suffix}.docx"
            report = apply_mode(contract, data, mode, out)
            print_report(report)
            any_failed = any_failed or bool(report["failed"])
        if any_failed:
            print("\n提示:存在定位失败条目。请核对 report-data.json 中的"
                  " original_text / suggestion_original 是否与合同逐字一致。")
    else:
        if not args.output:
            print("错误:单模式必须提供 --output", file=sys.stderr)
            sys.exit(1)
        out = Path(args.output)
        if out.suffix.lower() != ".docx":
            print(f"错误:--output 扩展名必须是 .docx,当前是 '{out.suffix}'",
                  file=sys.stderr)
            sys.exit(1)
        report = apply_mode(contract, data, args.mode, out)
        print_report(report)
        if report["failed"]:
            sys.exit(2)


if __name__ == "__main__":
    main()
