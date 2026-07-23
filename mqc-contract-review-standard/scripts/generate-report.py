#!/usr/bin/env python3
"""
generate-report.py · v1.0.0

生成合同审查报告 Word 文档

核心特性:
1. 严格对齐《参考格式》:16/14/12pt · 1.5行距 · 段前后精确 · 首缩0.85cm · 两端对齐
2. 全正文去色:所有"标题蓝"、"正文灰字"一律改黑;仅风险 ■ 符号保留功能色
3. 风险卡片三列对照表:条款原文 | 修改建议 | 修改理由
4. 修改建议内改动字用 #EE0000 正红(**...** 标记解析)
5. 规则来源 + 精准审查 + 缺失条款说明
6. _auto_compute_dimensions bug 修复
7. 封面原版小幅美化

作者:缪奇川
"""

import json
import sys
import argparse
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))

from _common import (
    get_brand_line,
    setup_page, setup_default_paragraph_normal,
    set_run_font, set_cell_shading, set_cell_borders, set_cell_vertical_center,
    setup_table_cell_paragraph, enforce_cell_formatting,
    add_heading_black, add_body_paragraph,
    apply_body_paragraph_format, apply_no_indent_paragraph_format,
    add_comparison_table, _add_revised_text_with_redmark,
    style_header_row_light, apply_data_row_style, style_data_table,
    FONT_CN_BODY, FONT_CN_TITLE, FONT_EN_ALL,
    FONT_SIZE_H1, FONT_SIZE_H2, FONT_SIZE_BODY, FONT_SIZE_SMALL, FONT_SIZE_MINI,
    LINE_SPACING,
    COLOR_RED, COLOR_AMBER, COLOR_GREEN, COLOR_BLACK, COLOR_REVISION_RED,
    HEX_TABLE_HEADER, HEX_TABLE_ZEBRA, HEX_DIVIDER_LINE,
    SYMBOL_RISK_SQUARE,
    RISK_LEVEL_COLOR, TIER_COLOR, TIER_DESC,
    CONFIRMATION_LEVEL_DESC, RULE_SOURCE_LABEL,
    SKILL_VERSION, FOOTER_COPYRIGHT, SKILL_SLOGAN,
)

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================
# 12 维度映射(修复 dimension bug)
# ============================================================

CHECKLIST_DIMENSIONS = [
    "一、主体资格",
    "二、合同标的与范围",
    "三、价款与支付",
    "四、权利与义务",
    "五、违约责任",
    "六、知识产权",
    "七、保密条款",
    "八、争议解决",
    "九、合同期限与终止",
    "十、不可抗力",
    "十一、送达与签署",
    "十二、其他常见条款",
]

CHECKLIST_ID_TO_DIMENSION = {
    **{f"C{i:03d}": 0 for i in range(1, 6)},    # 一、主体 C001-C005
    **{f"C{i:03d}": 1 for i in range(6, 10)},   # 二、标的 C006-C009
    **{f"C{i:03d}": 2 for i in range(10, 17)},  # 三、价款 C010-C016
    **{f"C{i:03d}": 3 for i in range(17, 21)},  # 四、权利义务 C017-C020
    **{f"C{i:03d}": 4 for i in range(21, 28)},  # 五、违约 C021-C027
    **{f"C{i:03d}": 5 for i in range(28, 32)},  # 六、知产 C028-C031
    **{f"C{i:03d}": 6 for i in range(32, 37)},  # 七、保密 C032-C036
    **{f"C{i:03d}": 7 for i in range(37, 42)},  # 八、争议 C037-C041
    **{f"C{i:03d}": 8 for i in range(42, 46)},  # 九、期限 C042-C045
    **{f"C{i:03d}": 9 for i in range(46, 50)},  # 十、不可抗力 C046-C049
    **{f"C{i:03d}": 10 for i in range(50, 54)}, # 十一、送达 C050-C053
    **{f"C{i:03d}": 11 for i in range(54, 61)}, # 十二、其他 C054-C060
}


# ============================================================
# 辅助函数
# ============================================================

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def get_risk_color(risk_level_text):
    if not risk_level_text:
        return COLOR_BLACK
    if "高风险" in risk_level_text or "🔴" in risk_level_text:
        return COLOR_RED
    if "中风险" in risk_level_text or "🟡" in risk_level_text:
        return COLOR_AMBER
    if "低风险" in risk_level_text or "🟢" in risk_level_text:
        return COLOR_GREEN
    return COLOR_BLACK


def normalize_risk_level(risk_level_text):
    if not risk_level_text:
        return ""
    if "高风险" in risk_level_text or "🔴" in risk_level_text:
        return f"{SYMBOL_RISK_SQUARE} 高风险"
    if "中风险" in risk_level_text or "🟡" in risk_level_text:
        return f"{SYMBOL_RISK_SQUARE} 中风险"
    if "低风险" in risk_level_text or "🟢" in risk_level_text:
        return f"{SYMBOL_RISK_SQUARE} 低风险"
    if "正常" in risk_level_text:
        return "✓ 正常"
    return risk_level_text


def add_divider_line(doc, color_hex=HEX_DIVIDER_LINE):
    """水平细分隔线(章节锚点)"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Cm(16)
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        b.set(qn("w:sz"), "0")
        tcBorders.append(b)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), color_hex)
    tcBorders.append(bottom)
    tcPr.append(tcBorders)
    cell.paragraphs[0].text = ""


def add_thin_line(doc, color_hex="1387C0", size_pt=8):
    """精致水平线(封面装饰用)"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Cm(8)
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        b.set(qn("w:sz"), "0")
        tcBorders.append(b)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size_pt))
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), color_hex)
    tcBorders.append(bottom)
    tcPr.append(tcBorders)
    cell.paragraphs[0].text = ""


def add_meta_line(doc, text, bold_prefix=None):
    """元数据行(不首缩)"""
    p = doc.add_paragraph()
    apply_no_indent_paragraph_format(p)
    if bold_prefix:
        run_pfx = p.add_run(bold_prefix)
        set_run_font(run_pfx, size=FONT_SIZE_BODY, bold=True, color=COLOR_BLACK)
    run = p.add_run(text)
    set_run_font(run, size=FONT_SIZE_BODY, color=COLOR_BLACK)
    return p


def setup_header_footer(section, data):
    """页眉页脚 · 黑字"""
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header_para.add_run(f"{data.get('contract_name', '')}  ·  合同审查报告")
    set_run_font(run, size=FONT_SIZE_MINI, color=COLOR_BLACK)

    footer = section.footer
    footer_para = footer.paragraphs[0]
    run = footer_para.add_run(f"审查编号:{data.get('review_number', '')}")
    set_run_font(run, size=FONT_SIZE_MINI, color=COLOR_BLACK)
    footer_para.add_run("\t\t")

    run_pn_label = footer_para.add_run("第 ")
    set_run_font(run_pn_label, size=FONT_SIZE_MINI, color=COLOR_BLACK)
    run_pn = footer_para.add_run()
    _add_field(run_pn, "PAGE")
    set_run_font(run_pn, size=FONT_SIZE_MINI, color=COLOR_BLACK)
    run_sep = footer_para.add_run(" 页 / 共 ")
    set_run_font(run_sep, size=FONT_SIZE_MINI, color=COLOR_BLACK)
    run_tot = footer_para.add_run()
    _add_field(run_tot, "NUMPAGES")
    set_run_font(run_tot, size=FONT_SIZE_MINI, color=COLOR_BLACK)
    run_end = footer_para.add_run(" 页")
    set_run_font(run_end, size=FONT_SIZE_MINI, color=COLOR_BLACK)


def _add_field(run, field_code):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)
    instr = OxmlElement("w:instrText")
    instr.text = field_code
    run._r.append(instr)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


# ============================================================
# 封面
# ============================================================

def render_cover(doc, data):
    """封面 · 包豪斯高端杂志风格 · 单外层表格保证单页

    设计语言:
    - 强网格·左对齐(包豪斯核心)
    - 浅蓝 #1387C0 单色 + 黑字
    - 英文大字全大写衬线 + 中文宋体加粗形成字重张力
    - 极简装饰:顶部色块 + 两条水平分隔线

    关键实现:
    - 把所有视觉元素放到一个外层 1 列 N 行表格
    - 每行一个元素,单元格内控制段落
    - 避免 python-docx 在每个独立 table 前插入隐藏段落导致溢出
    """
    # 外层大表格·1 列·每行承载一个视觉元素
    # 行数:15 (顶部色块/间距/英文3行/短横/间距/中文/长线/间距/信息表/间距/长线/间距/品牌/间距)
    # 但信息表本身是子表格,放到单元格里
    outer = doc.add_table(rows=14, cols=1)
    outer.alignment = WD_TABLE_ALIGNMENT.LEFT
    outer_cell_width = Cm(15)
    for row in outer.rows:
        row.cells[0].width = outer_cell_width

    row_idx = 0

    # 行 1:顶部浅蓝色块(精致3pt细条)
    _fill_block_cell(outer.cell(row_idx, 0), color_hex="1387C0",
                     width_cm=2.5, height_pt=3)
    row_idx += 1

    # 行 2:空白(约 30pt 高)
    _fill_spacer_cell(outer.cell(row_idx, 0), pt=30)
    row_idx += 1

    # 行 3:英文 CONTRACT(36pt TNR 加粗)
    _fill_text_cell(outer.cell(row_idx, 0), "CONTRACT",
                    size=36, bold=True, font_cn="宋体")
    row_idx += 1

    # 行 4:英文 REVIEW
    _fill_text_cell(outer.cell(row_idx, 0), "REVIEW",
                    size=36, bold=True, font_cn="宋体")
    row_idx += 1

    # 行 5:英文 REPORT
    _fill_text_cell(outer.cell(row_idx, 0), "REPORT",
                    size=36, bold=True, font_cn="宋体")
    row_idx += 1

    # 行 6:空白 10pt + 短横线 1.5pt(合并为一个单元格)
    _fill_spacer_cell(outer.cell(row_idx, 0), pt=8)
    row_idx += 1

    # 行 7:短横线(精致2pt)
    _fill_block_cell(outer.cell(row_idx, 0), color_hex="1387C0",
                     width_cm=2.5, height_pt=2)
    row_idx += 1

    # 行 8:空白
    _fill_spacer_cell(outer.cell(row_idx, 0), pt=12)
    row_idx += 1

    # 行 9:中文标题 22pt 宋体加粗
    _fill_text_cell(outer.cell(row_idx, 0), "合同审查报告",
                    size=22, bold=True, font_cn="宋体")
    row_idx += 1

    # 行 10:长分隔线(精致1pt细线)
    _fill_block_cell(outer.cell(row_idx, 0), color_hex="1387C0",
                     width_cm=14, height_pt=1)
    row_idx += 1

    # 行 11:信息表(作为嵌套表格放到单元格内)
    _fill_info_table_cell(outer.cell(row_idx, 0), data)
    row_idx += 1

    # 行 12:长分隔线(底·精致1pt细线)
    _fill_block_cell(outer.cell(row_idx, 0), color_hex="1387C0",
                     width_cm=14, height_pt=1)
    row_idx += 1

    # 行 13:间距
    _fill_spacer_cell(outer.cell(row_idx, 0), pt=6)
    row_idx += 1

    # 行 14:品牌署名
    _fill_text_cell(outer.cell(row_idx, 0),
        get_brand_line(data) or "",
        size=FONT_SIZE_MINI, bold=False, font_cn="宋体")
    row_idx += 1

    # 外层表格所有单元格无边框、无缩进
    _clean_outer_table_cells(outer)

    # 分页符
    add_page_break(doc)


def _clean_outer_table_cells(table):
    """封面外层表格:所有单元格去边框"""
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._element.get_or_add_tcPr()
            # 去边框
            existing = tcPr.find(qn("w:tcBorders"))
            if existing is not None:
                tcPr.remove(existing)
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "right", "bottom"):
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"), "none")
                b.set(qn("w:sz"), "0")
                tcBorders.append(b)
            tcPr.append(tcBorders)
            # 去单元格内边距(避免多余空间)
            tcMar = OxmlElement("w:tcMar")
            for side in ("top", "left", "right", "bottom"):
                m = OxmlElement(f"w:{side}")
                m.set(qn("w:w"), "0")
                m.set(qn("w:type"), "dxa")
                tcMar.append(m)
            # 检查已有 tcMar
            existing_mar = tcPr.find(qn("w:tcMar"))
            if existing_mar is not None:
                tcPr.remove(existing_mar)
            tcPr.append(tcMar)


def _fill_block_cell(cell, color_hex="1387C0", width_cm=2.5, height_pt=2):
    """封面用:在外层单元格中放一个色块(用嵌套小表格实现)
    height_pt: 色块目标高度(pt) · 通过行高属性精确控制(不依赖字号)
    """
    # 清空该外层 cell 默认段落
    p_existing = cell.paragraphs[0]
    p_existing.paragraph_format.first_line_indent = Cm(0)
    p_existing.paragraph_format.space_before = Pt(0)
    p_existing.paragraph_format.space_after = Pt(0)
    p_existing.paragraph_format.line_spacing = 1.0
    # 在其后追加嵌套小色块表
    inner_table = cell.add_table(rows=1, cols=1)
    inner_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    ic = inner_table.cell(0, 0)
    ic.width = Cm(width_cm)
    set_cell_shading(ic, color_hex)
    # 去边框
    tcPr = ic._element.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "right", "bottom"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        b.set(qn("w:sz"), "0")
        tcBorders.append(b)
    tcPr.append(tcBorders)
    # 去单元格内边距(上下左右都设为 0)
    tcMar = OxmlElement("w:tcMar")
    for side in ("top", "left", "right", "bottom"):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), "0")
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)
    # 段落设为最小
    p_in = ic.paragraphs[0]
    p_in.paragraph_format.first_line_indent = Cm(0)
    p_in.paragraph_format.space_before = Pt(0)
    p_in.paragraph_format.space_after = Pt(0)
    p_in.paragraph_format.line_spacing = Pt(1)  # 绝对 1pt 行距
    p_in.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    # 不加任何 run,段落空的,但有色块底色
    # 精确控制行高:用 trH exact 模式
    tr = inner_table.rows[0]._element
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = OxmlElement("w:trPr")
        tr.insert(0, trPr)
    # 行高用 twip 单位(1pt = 20 twip)
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(height_pt * 20)))
    trHeight.set(qn("w:hRule"), "exact")
    trPr.append(trHeight)
    # 占位段落字号最小
    occupant = p_existing.add_run("")
    set_run_font(occupant, size=1, color=COLOR_BLACK)


def _fill_spacer_cell(cell, pt=10):
    """封面用:纯空白间距 cell"""
    p = cell.paragraphs[0]
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(" ")
    set_run_font(run, size=pt, color=COLOR_BLACK)


def _fill_text_cell(cell, text, size=12, bold=False, font_cn="宋体"):
    """封面用:在外层单元格放一段左对齐文字"""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    run = p.add_run(text)
    set_run_font(run, font_name_cn=font_cn, size=size, bold=bold, color=COLOR_BLACK)


def _fill_info_table_cell(cell, data):
    """封面用:在外层单元格放一个信息子表"""
    # 先保证外层 cell 原段落不占空间
    p_existing = cell.paragraphs[0]
    p_existing.paragraph_format.first_line_indent = Cm(0)
    p_existing.paragraph_format.space_before = Pt(0)
    p_existing.paragraph_format.space_after = Pt(0)
    p_existing.paragraph_format.line_spacing = 1.0
    r = p_existing.add_run("")
    set_run_font(r, size=1, color=COLOR_BLACK)

    cover_items = [
        ("合同名称", data.get("contract_name", "")),
        ("审查编号", data.get("review_number", "")),
        ("委托方", data.get("client", "")),
        ("经办律师", " · ".join(x for x in (
            str(data.get("lawyer", "")).strip(),
            str(data.get("law_firm", "")).strip()) if x)),
        ("审查立场", data.get("review_stance", "")),
        ("审查日期", data.get("review_date", "")),
    ]
    cover_items_f = [(l, v) for l, v in cover_items if v]

    inner = cell.add_table(rows=len(cover_items_f), cols=2)
    inner.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (label, value) in enumerate(cover_items_f):
        c0 = inner.cell(i, 0)
        c1 = inner.cell(i, 1)
        c0.width = Cm(3.2)
        c1.width = Cm(11)
        set_cell_vertical_center(c0)
        set_cell_vertical_center(c1)
        # 去边框
        for cc in (c0, c1):
            tcPr = cc._element.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "right", "bottom"):
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"), "none")
                b.set(qn("w:sz"), "0")
                tcBorders.append(b)
            tcPr.append(tcBorders)

        p0 = c0.paragraphs[0]
        p0.paragraph_format.first_line_indent = Cm(0)
        p0.paragraph_format.space_before = Pt(1)
        p0.paragraph_format.space_after = Pt(1)
        p0.paragraph_format.line_spacing = 1.1
        run = p0.add_run(label)
        set_run_font(run, size=FONT_SIZE_BODY, color=COLOR_BLACK)

        p1 = c1.paragraphs[0]
        p1.paragraph_format.first_line_indent = Cm(0)
        p1.paragraph_format.space_before = Pt(1)
        p1.paragraph_format.space_after = Pt(1)
        p1.paragraph_format.line_spacing = 1.1
        run = p1.add_run(value)
        set_run_font(run, size=FONT_SIZE_BODY, bold=True, color=COLOR_BLACK)


def render_basic_info(doc, data):
    add_heading_black(doc, "一、审查基本信息", level=1)

    info_items = [
        ("合同名称", data.get("contract_name", "")),
        ("合同金额", data.get("amount", "")),
        ("合同期限", data.get("term", "")),
        ("甲方", data.get("party_a", "")),
        ("乙方", data.get("party_b", "")),
        ("审查立场", data.get("review_stance", "")),
        ("审查日期", data.get("review_date", "")),
        ("审查依据", f"标准审查清单 v{data.get('checklist_version', '1.0.0')}"),
        ("审查范围", data.get("scope", "合同正文 + 附件")),
        ("审查编号", data.get("review_number", "")),
    ]

    table = doc.add_table(rows=len(info_items), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (label, value) in enumerate(info_items):
        c0 = table.cell(i, 0)
        c1 = table.cell(i, 1)
        c0.width = Cm(3.5)
        c1.width = Cm(13.5)
        set_cell_vertical_center(c0)
        set_cell_vertical_center(c1)
        set_cell_borders(c0)
        set_cell_borders(c1)
        p0 = c0.paragraphs[0]
        p0.paragraph_format.first_line_indent = None
        run = p0.add_run(label)
        set_run_font(run, size=FONT_SIZE_BODY, bold=True, color=COLOR_BLACK)
        p1 = c1.paragraphs[0]
        p1.paragraph_format.first_line_indent = None
        run = p1.add_run(value)
        set_run_font(run, size=FONT_SIZE_BODY, color=COLOR_BLACK)
    enforce_cell_formatting(table, space_pt=3)


# ============================================================
# 二、执行摘要
# ============================================================

def render_executive_summary(doc, data):
    add_heading_black(doc, "二、执行摘要", level=1)
    add_body_paragraph(doc, "本章以一页篇幅呈现报告的核心结论,便于决策者快速了解。")

    add_heading_black(doc, "2.1 综合风险评级", level=2)
    rating = data.get("overall_rating", "")
    rating_flat = normalize_risk_level(rating) or rating
    rating_color = get_risk_color(rating)
    p = doc.add_paragraph()
    apply_no_indent_paragraph_format(p)
    run = p.add_run("评级:  ")
    set_run_font(run, size=FONT_SIZE_BODY, bold=True, color=COLOR_BLACK)
    if rating_flat:
        run = p.add_run(rating_flat)
        set_run_font(run, size=FONT_SIZE_BODY, bold=True, color=rating_color)

    add_heading_black(doc, "2.2 可签性判断", level=2)
    signability = data.get("signability", "")
    sign_flat = signability
    for emoji in ("🔴", "🟡", "🟢"):
        sign_flat = sign_flat.replace(emoji, SYMBOL_RISK_SQUARE)
    sign_color = get_risk_color(signability)
    p = doc.add_paragraph()
    apply_no_indent_paragraph_format(p)
    run = p.add_run("判断:  ")
    set_run_font(run, size=FONT_SIZE_BODY, bold=True, color=COLOR_BLACK)
    if sign_flat:
        run = p.add_run(sign_flat)
        set_run_font(run, size=FONT_SIZE_BODY, bold=True, color=sign_color)

    add_heading_black(doc, "2.3 审查结论摘要", level=2)
    summary_text = data.get("summary", "")
    for emoji in ("🔴", "🟡", "🟢"):
        summary_text = summary_text.replace(emoji, SYMBOL_RISK_SQUARE)
    add_body_paragraph(doc, summary_text)

    add_heading_black(doc, "2.4 优先处理事项(Top 3)", level=2)
    top_3 = data.get("top_3", [])
    if top_3:
        table = doc.add_table(rows=1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["序号", "问题", "风险等级", "对应清单", "建议行动"]
        # 列宽调整:问题列 3.5 建议行动列 6.3
        col_widths = [Cm(1.0), Cm(3.8), Cm(2.0), Cm(1.8), Cm(6.2)]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.width = col_widths[i]
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = None
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h)
            set_run_font(run, size=FONT_SIZE_MINI, bold=True, color=COLOR_BLACK)

        for item in top_3:
            row = table.add_row()
            risk = item.get("risk_level", "")
            risk_flat = normalize_risk_level(risk)
            risk_color = get_risk_color(risk)
            vals = [
                str(item.get("seq", "")),
                item.get("issue", ""),
                risk_flat,
                item.get("checklist_id", ""),
                item.get("action", ""),
            ]
            for i, v in enumerate(vals):
                cell = row.cells[i]
                cell.width = col_widths[i]
                p = cell.paragraphs[0]
                p.paragraph_format.first_line_indent = None
                if i in (0, 2, 3):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(v)
                if i == 2 and risk_flat:
                    set_run_font(run, size=FONT_SIZE_MINI, bold=True, color=risk_color)
                else:
                    set_run_font(run, size=FONT_SIZE_MINI, color=COLOR_BLACK)

        style_data_table(table, has_header=True)
        enforce_cell_formatting(table, space_pt=2)
    else:
        add_body_paragraph(doc, "无高风险事项。")


# ============================================================
# 三、风险发现汇总(新增"来源"列 · 3.3 规则来源分布节 · 维度真实映射)
# ============================================================

def render_findings_summary(doc, data):
    add_heading_black(doc, "三、风险发现汇总", level=1)

    findings = data.get("findings", [])

    # 3.1 汇总表(新增规则来源列)
    add_heading_black(doc, "3.1 汇总表", level=2)
    if findings:
        table = doc.add_table(rows=1, cols=7)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["序号", "风险等级", "层级", "条款位置", "清单", "来源", "问题概述"]
        # 列宽调整+字号调小:避免"第12.2条"换行
        col_widths = [Cm(1.0), Cm(1.7), Cm(1.1), Cm(2.0), Cm(1.2), Cm(1.8), Cm(7.0)]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.width = col_widths[i]
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = None
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h)
            set_run_font(run, size=FONT_SIZE_MINI, bold=True, color=COLOR_BLACK)

        for idx, finding in enumerate(findings, 1):
            row = table.add_row()
            risk = finding.get("risk_level", "")
            risk_flat = normalize_risk_level(risk)
            risk_color = get_risk_color(risk)
            rule_source = finding.get("rule_source", "checklist")
            source_label = RULE_SOURCE_LABEL.get(rule_source, rule_source)
            is_precise = finding.get("is_precise_review", False)
            summary_text = finding.get("issue_summary", finding.get("title", ""))
            if is_precise:
                summary_text = f"[精准] {summary_text}"

            vals = [
                str(idx), risk_flat,
                finding.get("confirmation_level", ""),
                finding.get("clause_location", ""),
                finding.get("checklist_id", ""),
                source_label, summary_text,
            ]
            for i, v in enumerate(vals):
                cell = row.cells[i]
                cell.width = col_widths[i]
                p = cell.paragraphs[0]
                p.paragraph_format.first_line_indent = None
                if i in (0, 1, 2, 4, 5):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(v)
                if i == 1 and risk_flat:
                    set_run_font(run, size=FONT_SIZE_MINI, bold=True, color=risk_color)
                else:
                    set_run_font(run, size=FONT_SIZE_MINI, color=COLOR_BLACK)

        style_data_table(table, has_header=True)
        enforce_cell_formatting(table, space_pt=2)
    else:
        add_body_paragraph(doc, "本次审查未发现风险项。")

    # 3.2 风险等级统计
    add_heading_black(doc, "3.2 风险等级统计", level=2)
    red_count = sum(1 for f in findings if "高风险" in f.get("risk_level", "") or f.get("risk_level", "").startswith("🔴"))
    amber_count = sum(1 for f in findings if "中风险" in f.get("risk_level", "") or f.get("risk_level", "").startswith("🟡"))
    green_count = sum(1 for f in findings if "低风险" in f.get("risk_level", "") or f.get("risk_level", "").startswith("🟢"))
    l1_count = sum(1 for f in findings if f.get("confirmation_level", "").startswith("L1"))
    l2_count = sum(1 for f in findings if f.get("confirmation_level", "").startswith("L2"))
    l3_count = sum(1 for f in findings if f.get("confirmation_level", "").startswith("L3"))

    add_body_paragraph(doc, f"本合同共发现 {len(findings)} 处风险:")

    p = doc.add_paragraph()
    apply_body_paragraph_format(p)
    run = p.add_run(SYMBOL_RISK_SQUARE + " ")
    set_run_font(run, size=FONT_SIZE_BODY, bold=True, color=COLOR_RED)
    run = p.add_run(f"高风险:{red_count} 处     ")
    set_run_font(run, size=FONT_SIZE_BODY, color=COLOR_BLACK)
    run = p.add_run(SYMBOL_RISK_SQUARE + " ")
    set_run_font(run, size=FONT_SIZE_BODY, bold=True, color=COLOR_AMBER)
    run = p.add_run(f"中风险:{amber_count} 处     ")
    set_run_font(run, size=FONT_SIZE_BODY, color=COLOR_BLACK)
    run = p.add_run(SYMBOL_RISK_SQUARE + " ")
    set_run_font(run, size=FONT_SIZE_BODY, bold=True, color=COLOR_GREEN)
    run = p.add_run(f"低风险:{green_count} 处")
    set_run_font(run, size=FONT_SIZE_BODY, color=COLOR_BLACK)

    add_body_paragraph(doc, "按确认层级分布:")
    for line in [
        f"L1 基础级(AI 可独立确认):{l1_count} 处",
        f"L2 中级(建议人工复核):{l2_count} 处",
        f"L3 高级(必须人工复核):{l3_count} 处",
    ]:
        add_body_paragraph(doc, line)

    # 3.3 规则来源分布
    add_heading_black(doc, "3.3 规则来源分布", level=2)
    source_counts = {"checklist": 0, "playbook": 0, "extra": 0}
    for f in findings:
        src = f.get("rule_source", "checklist")
        source_counts[src] = source_counts.get(src, 0) + 1

    add_body_paragraph(doc,
        "按来源归类,区分:"
        "标准清单触发(基于 checklist.md)、"
        "Playbook 覆盖(组织立场手册)、"
        "清单外发现(AI 扩展判断)。")

    for src_key in ("checklist", "playbook", "extra"):
        add_body_paragraph(doc,
            f"{RULE_SOURCE_LABEL[src_key]}:{source_counts.get(src_key, 0)} 处")

    # 3.4 风险分布维度图(真实映射 · 修复 bug · 仅显示非零维度)
    add_heading_black(doc, "3.4 风险分布维度图", level=2)
    dim_stats = _auto_compute_dimensions(findings)
    non_zero = [d for d in dim_stats if d["count"] > 0]

    if not non_zero:
        add_body_paragraph(doc, "12 个维度均未发现风险项。")
        return

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["维度", "风险项数", "涉及清单 ID"]
    col_widths = [Cm(5), Cm(2.5), Cm(9)]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = col_widths[i]
        p = cell.paragraphs[0]
        p.paragraph_format.first_line_indent = None
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, size=FONT_SIZE_SMALL, bold=True, color=COLOR_BLACK)

    for item in non_zero:
        row = table.add_row()
        vals = [item["dimension"], str(item["count"]), " · ".join(item["checklist_ids"])]
        for i, v in enumerate(vals):
            cell = row.cells[i]
            cell.width = col_widths[i]
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = None
            if i == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(v)
            set_run_font(run, size=FONT_SIZE_SMALL, color=COLOR_BLACK)

    style_data_table(table, has_header=True)
    enforce_cell_formatting(table)


def _auto_compute_dimensions(findings):
    stats = [{"dimension": d, "count": 0, "checklist_ids": []} for d in CHECKLIST_DIMENSIONS]
    for finding in findings:
        cid = finding.get("checklist_id", "")
        if cid in CHECKLIST_ID_TO_DIMENSION:
            idx = CHECKLIST_ID_TO_DIMENSION[cid]
            stats[idx]["count"] += 1
            if cid not in stats[idx]["checklist_ids"]:
                stats[idx]["checklist_ids"].append(cid)
    return stats


# ============================================================
# 四、逐条审查详情(三列对照表)
# ============================================================

def render_findings_detail(doc, data):
    add_heading_black(doc, "四、逐条审查详情", level=1)
    findings = data.get("findings", [])
    if not findings:
        add_body_paragraph(doc, "本次审查未发现风险项。")
        return
    for idx, finding in enumerate(findings, 1):
        _render_single_finding(doc, idx, finding)


def _render_single_finding(doc, idx, finding):
    """风险卡片
    结构:
      1. 标题行:"风险 XX  ■ 高风险  [精准审查]  标题"
      2. 元数据行:条款位置 · 对应清单 · 确认层级 · 规则来源 · 谈判优先级
      3. 触发标准行 / 立场判断行
      4. 三列对照表:条款原文 | 修改建议 | 修改理由
      5. 相关法条行
      6. 分隔线
    """
    risk_level = finding.get("risk_level", "")
    risk_color = get_risk_color(risk_level)
    risk_flat = normalize_risk_level(risk_level)
    title_text = finding.get("title", "")
    is_precise = finding.get("is_precise_review", False)

    # 1. 标题
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(16)
    heading.paragraph_format.space_after = Pt(4)
    heading.paragraph_format.first_line_indent = None

    run = heading.add_run(f"风险 {idx:02d}   ")
    set_run_font(run, size=FONT_SIZE_H2, bold=True, color=COLOR_BLACK)

    if risk_flat:
        run = heading.add_run(risk_flat + "   ")
        set_run_font(run, size=FONT_SIZE_H2, bold=True, color=risk_color)

    if is_precise:
        run = heading.add_run("[精准审查]   ")
        set_run_font(run, size=FONT_SIZE_H2, bold=True, color=COLOR_BLACK)

    run = heading.add_run(title_text)
    set_run_font(run, size=FONT_SIZE_H2, bold=True, color=COLOR_BLACK)

    # 2. 元数据行
    rule_source = finding.get("rule_source", "checklist")
    source_label = RULE_SOURCE_LABEL.get(rule_source, rule_source)
    negotiation_tier = finding.get("negotiation_tier", "")

    meta_items = []
    if finding.get("clause_location"):
        meta_items.append(f"条款位置:{finding['clause_location']}")
    if finding.get("checklist_id"):
        meta_items.append(f"对应清单:{finding['checklist_id']}")
    if finding.get("confirmation_level"):
        meta_items.append(f"确认层级:{finding['confirmation_level']}")
    meta_items.append(f"规则来源:{source_label}")
    if negotiation_tier:
        meta_items.append(f"谈判优先级:{negotiation_tier}")

    p = doc.add_paragraph()
    apply_no_indent_paragraph_format(p)
    run = p.add_run("  ·  ".join(meta_items))
    set_run_font(run, size=FONT_SIZE_BODY, color=COLOR_BLACK)

    # 3. 触发标准
    trigger = finding.get("trigger_standard", "")
    for emoji in ("🔴", "🟡", "🟢"):
        trigger = trigger.replace(emoji, SYMBOL_RISK_SQUARE)
    if trigger:
        p = doc.add_paragraph()
        apply_no_indent_paragraph_format(p)
        run_pfx = p.add_run("触发标准:")
        set_run_font(run_pfx, size=FONT_SIZE_BODY, bold=True, color=COLOR_BLACK)
        run = p.add_run(trigger)
        set_run_font(run, size=FONT_SIZE_BODY, color=COLOR_BLACK)

    # 立场判断
    position_diff = finding.get("position_difference", "")
    if position_diff and position_diff.strip() != "中性":
        p = doc.add_paragraph()
        apply_no_indent_paragraph_format(p)
        run_pfx = p.add_run("立场判断:")
        set_run_font(run_pfx, size=FONT_SIZE_BODY, bold=True, color=COLOR_BLACK)
        run = p.add_run(position_diff)
        set_run_font(run, size=FONT_SIZE_BODY, color=COLOR_BLACK)

    # 4. 三列对照表(核心)
    original_text = finding.get("original_text", "") or finding.get("suggestion_original", "")
    revised_text = finding.get("suggestion_revised", "") or finding.get("suggestion", "")
    reason_text = finding.get("risk_description", "") or finding.get("reason", "")

    if original_text and revised_text:
        p_spacer = doc.add_paragraph()
        p_spacer.paragraph_format.space_after = Pt(0)
        p_spacer.paragraph_format.first_line_indent = None
        add_comparison_table(doc,
            original_text=original_text,
            revised_text=revised_text,
            reason_text=reason_text,
            reason_label="修改理由",
        )
    elif revised_text:
        p = doc.add_paragraph()
        apply_no_indent_paragraph_format(p)
        run_pfx = p.add_run("修改建议:")
        set_run_font(run_pfx, size=FONT_SIZE_BODY, bold=True, color=COLOR_BLACK)
        p = doc.add_paragraph()
        apply_body_paragraph_format(p)
        _add_revised_text_with_redmark(p, revised_text)

    # 5. 相关法条
    laws = finding.get("related_laws", "")
    if laws:
        p = doc.add_paragraph()
        apply_no_indent_paragraph_format(p)
        p.paragraph_format.space_before = Pt(4)
        run_pfx = p.add_run("相关法条:")
        set_run_font(run_pfx, size=FONT_SIZE_BODY, bold=True, color=COLOR_BLACK)
        run = p.add_run(laws)
        set_run_font(run, size=FONT_SIZE_BODY, color=COLOR_BLACK)

    # 6. 分隔线
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    add_divider_line(doc)


# ============================================================
# 五、缺失条款提示(显式说明机制)
# ============================================================

def render_missing_clauses(doc, data):
    missing = data.get("missing_clauses", [])
    if not missing:
        return

    add_heading_black(doc, "五、缺失条款提示", level=1)

    add_body_paragraph(doc,
        "本章列出审查清单中标记为\"缺失检查项\"(is_completion_check=true)、"
        "但合同文本未发现对应条款的项目。"
        "清单外条款不在此章列出。")

    # 5.1 汇总表
    add_heading_black(doc, "5.1 缺失条款汇总", level=2)
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["序号", "缺失条款", "对应清单", "重要程度", "建议概述"]
    col_widths = [Cm(1.2), Cm(4), Cm(2), Cm(2.5), Cm(6.5)]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = col_widths[i]
        p = cell.paragraphs[0]
        p.paragraph_format.first_line_indent = None
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, size=FONT_SIZE_SMALL, bold=True, color=COLOR_BLACK)

    for idx, item in enumerate(missing, 1):
        row = table.add_row()
        importance_raw = item.get("importance", "")
        importance_clean = importance_raw
        for emoji in ("🔴", "🟡", "🟢"):
            importance_clean = importance_clean.replace(emoji, SYMBOL_RISK_SQUARE)
        if "必须" in importance_raw:
            importance_color = COLOR_RED
        elif "建议" in importance_raw:
            importance_color = COLOR_AMBER
        else:
            importance_color = COLOR_BLACK

        vals = [
            str(idx),
            item.get("clause", ""),
            item.get("checklist_id", ""),
            importance_clean,
            item.get("suggestion_brief", item.get("suggestion", ""))[:60],
        ]
        for i, v in enumerate(vals):
            cell = row.cells[i]
            cell.width = col_widths[i]
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = None
            if i in (0, 2, 3):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(v)
            if i == 3 and v:
                set_run_font(run, size=FONT_SIZE_SMALL, bold=True, color=importance_color)
            else:
                set_run_font(run, size=FONT_SIZE_SMALL, color=COLOR_BLACK)

    style_data_table(table, has_header=True)
    enforce_cell_formatting(table)

    # 5.2 缺失条款补充建议(三列对照表 · 补充理由)
    add_heading_black(doc, "5.2 缺失条款补充建议", level=2)
    add_body_paragraph(doc,
        "对每项\"必须补充\"和\"建议补充\"的缺失条款,"
        "给出可直接插入合同的桥面条款。")

    for idx, item in enumerate(missing, 1):
        bridge = item.get("bridge_clause", "")
        if not bridge:
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.first_line_indent = None
        run = p.add_run(
            f"【缺失 {idx}】{item.get('clause', '')}"
            f"  ·  对应清单:{item.get('checklist_id', '')}"
            f"  ·  {item.get('importance', '')}"
        )
        set_run_font(run, size=FONT_SIZE_BODY, bold=True, color=COLOR_BLACK)

        add_comparison_table(doc,
            original_text="(合同中未约定此条款)",
            revised_text=bridge,
            reason_text=item.get("reason",
                item.get("suggestion",
                    f"本条款是{item.get('clause', '')}的标准条款,"
                    "合同中缺失该条款将影响相关权利义务的明确性。")),
            reason_label="补充理由",
        )

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        add_divider_line(doc)


# ============================================================
# 六、清单外发现
# ============================================================

def render_extra_findings(doc, data):
    extra = data.get("extra_findings", [])
    if not extra:
        return

    add_heading_black(doc, "六、清单外发现", level=1)
    add_body_paragraph(doc,
        "审查过程中发现的、当前审查清单未覆盖但值得关注的问题。"
        "这是 checklist.md 持续迭代的核心输入。")

    for idx, item in enumerate(extra, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.first_line_indent = None
        run = p.add_run(f"【清单外发现 {idx}】{item.get('title', '')}")
        set_run_font(run, size=FONT_SIZE_BODY, bold=True, color=COLOR_BLACK)

        if item.get("clause_location"):
            add_meta_line(doc, item['clause_location'], bold_prefix="条款位置:")

        if item.get("original_text"):
            p = doc.add_paragraph()
            apply_no_indent_paragraph_format(p)
            run_pfx = p.add_run("合同原文:")
            set_run_font(run_pfx, size=FONT_SIZE_BODY, bold=True, color=COLOR_BLACK)
            run = p.add_run(item['original_text'])
            set_run_font(run, size=FONT_SIZE_BODY, color=COLOR_BLACK)

        if item.get("analysis"):
            p = doc.add_paragraph()
            apply_body_paragraph_format(p)
            run = p.add_run(item['analysis'])
            set_run_font(run, size=FONT_SIZE_BODY, color=COLOR_BLACK)

        if item.get("suggestion"):
            p = doc.add_paragraph()
            apply_no_indent_paragraph_format(p)
            run_pfx = p.add_run("建议处理:")
            set_run_font(run_pfx, size=FONT_SIZE_BODY, bold=True, color=COLOR_BLACK)
            run = p.add_run(item['suggestion'])
            set_run_font(run, size=FONT_SIZE_BODY, color=COLOR_BLACK)

        if item.get("suggest_add_to"):
            p = doc.add_paragraph()
            apply_no_indent_paragraph_format(p)
            run_pfx = p.add_run("建议补充到清单:")
            set_run_font(run_pfx, size=FONT_SIZE_BODY, bold=True, color=COLOR_BLACK)
            run = p.add_run(item['suggest_add_to'])
            set_run_font(run, size=FONT_SIZE_BODY, color=COLOR_BLACK)

        add_divider_line(doc)


# ============================================================
# 七、法条引用索引
# ============================================================

def render_laws_index(doc, data):
    laws = data.get("laws_index", [])
    if not laws:
        return

    add_heading_black(doc, "七、法条引用索引", level=1)
    add_body_paragraph(doc,
        "本章集中列出报告中引用的全部法律法规和条文。"
        "每条引用包含:法规全称、版本年份、条文编号、原文引用、校验日期。")

    for item in laws:
        pending = item.get("pending_verify", False)

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.first_line_indent = None
        run = p.add_run(f"【{item.get('id', '')}】 ")
        set_run_font(run, size=FONT_SIZE_BODY, bold=True, color=COLOR_BLACK)
        if pending:
            run = p.add_run("[待核实]  ")
            set_run_font(run, size=FONT_SIZE_BODY, bold=True, color=COLOR_AMBER)
        run = p.add_run(item.get("law_full_name", ""))
        set_run_font(run, size=FONT_SIZE_BODY, bold=True, color=COLOR_BLACK)
        if item.get("version_year"):
            run = p.add_run(f"({item['version_year']})")
            set_run_font(run, size=FONT_SIZE_BODY, color=COLOR_BLACK)
        if item.get("article"):
            run = p.add_run(f" {item['article']}")
            set_run_font(run, size=FONT_SIZE_BODY, bold=True, color=COLOR_BLACK)

        if item.get("original_text"):
            p = doc.add_paragraph()
            apply_body_paragraph_format(p)
            run_pfx = p.add_run("原文:")
            set_run_font(run_pfx, size=FONT_SIZE_BODY, bold=True, color=COLOR_BLACK)
            run = p.add_run(item['original_text'])
            set_run_font(run, size=FONT_SIZE_BODY, color=COLOR_BLACK)

        cited = item.get("cited_in", "")
        verify = item.get("verify_date", "")
        if cited or verify:
            p = doc.add_paragraph()
            apply_no_indent_paragraph_format(p)
            run = p.add_run(f"引用于:{cited}    校验日期:{verify}")
            set_run_font(run, size=FONT_SIZE_SMALL, color=COLOR_BLACK)


# ============================================================
# 八、后续操作建议
# ============================================================

def render_cross_skill_nav(doc, data):
    add_heading_black(doc, "八、后续操作建议", level=1)
    add_body_paragraph(doc, "根据本次审查结果,可调用以下同系列 Skill 协同工作:")

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["需求场景", "对应 Skill"]
    col_widths = [Cm(8), Cm(8)]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = col_widths[i]
        p = cell.paragraphs[0]
        p.paragraph_format.first_line_indent = None
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, size=FONT_SIZE_SMALL, bold=True, color=COLOR_BLACK)

    nav_items = [
        ("应用修改建议生成修订版合同", "mqc-contract-review-to-modify"),
        ("生成批注版合同(原文不改,仅加批注)", "mqc-contract-annotate"),
        ("查询合同主体的工商登记、经营状况", "mqc-entity-verify"),
        ("合同版本差异对比", "mqc-contract-template-review"),
        ("快速判断合同是否可签(3 分钟对话式)", "mqc-contract-review-quick"),
        ("批量审查多份合同", "mqc-contract-review-batch"),
        ("合同文件脱敏处理", "mqc-doc-sanitize"),
    ]

    for need, skill in nav_items:
        row = table.add_row()
        p = row.cells[0].paragraphs[0]
        p.paragraph_format.first_line_indent = None
        run = p.add_run(need)
        set_run_font(run, size=FONT_SIZE_SMALL, color=COLOR_BLACK)
        p = row.cells[1].paragraphs[0]
        p.paragraph_format.first_line_indent = None
        run = p.add_run(skill)
        set_run_font(run, size=FONT_SIZE_SMALL, bold=True, color=COLOR_BLACK)

    style_data_table(table, has_header=True)
    enforce_cell_formatting(table)


# ============================================================
# 九、重要声明
# ============================================================

def render_disclaimer(doc, data):
    add_heading_black(doc, "九、重要声明", level=1)

    declarations = [
        ("1. AI 辅助性质",
         "本审查报告由 AI 工具辅助生成,基于作者定义的审查清单逐条审查。"
         "所有风险分析和修改建议仅供参考,必须由执业律师审核确认后方可作为法律意见使用。"
         "本报告不构成法律意见书。"),
        ("2. 审查时效性",
         "本报告基于审查日期时的合同文本及当时有效的法律法规。"
         "合同文本如有变更(包括附件、补充协议、口头补充约定),需重新审查。"
         "法律法规后续变动可能影响本报告的部分结论。"),
        ("3. 审查范围限制",
         "本报告仅限于对合同文本本身的审查,不包括合同主体的工商登记、"
         "合同背景事实的真实性核实、合同商业条件的合理性评估、"
         "合同履行中的实际风险等外部事项。"),
        ("4. 清单覆盖边界",
         f"本报告基于审查清单 v{data.get('checklist_version', '1.0.0')} 逐条审查。"
         "清单未覆盖的风险可能未被发现。如发现清单外风险,"
         "已在第六章\"清单外发现\"列出。"),
        ("5. 法条引用说明",
         "本报告引用的法条已尽力核验,但 AI 工具对法条的记忆可能存在偏差。"
         "所有法条引用以该法律的官方发布版本为准。"
         "标注[待核实]的法条请律师特别复核。"),
        ("6. 使用范围限制",
         "本报告仅供本次审查的委托方及其指定人员使用,"
         "未经作者书面同意,不得对外披露、转载或用于除审查委托目的之外的其他用途。"),
        ("7. 律师独立判断",
         "本报告的任何结论和建议不替代执业律师在具体案件中的独立判断。"
         "律师应根据本次交易的具体背景、当事人的商业诉求、过往交易习惯等综合判断。"),
        ("8. 合同文本完整性说明",
         "如合同文本中包含 AI 无法识别的信息(如手写批注、图片签章、扫描页面的模糊文字),"
         "可能影响本报告的完整性。审查前应尽量使用可机读的电子版合同文本。"),
    ]

    for title_line, body_text in declarations:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.first_line_indent = None
        run = p.add_run(title_line)
        set_run_font(run, size=FONT_SIZE_BODY, bold=True, color=COLOR_BLACK)
        add_body_paragraph(doc, body_text)


# ============================================================
# 十、附录
# ============================================================

def render_appendix(doc, data):
    add_heading_black(doc, "十、附录", level=1)

    add_heading_black(doc, "附录 A · 审查方法论", level=2)
    add_body_paragraph(doc, "本审查依据以下原则执行:")
    for principle in [
        "(1) 分层结构:按 12 维度分组,逐条检查",
        "(2) 数字化判断标准:每项检查带明确阈值",
        "(3) 风险等级预设:清单内已标注默认风险等级,AI 执行而非自行判断",
        "(4) 合同类型适配:本审查基于标准版清单,如有专用清单可切换",
        "(5) 规则来源透明:每条发现标注来源(标准清单/Playbook/清单外)",
        "(6) 精准审查:对关键条款独立重扫一次,对抗大模型注意力涣散",
    ]:
        add_body_paragraph(doc, principle)

    add_heading_black(doc, "附录 B · 审查清单版本信息", level=2)
    for label, value in [
        ("清单版本", f"v{data.get('checklist_version', '1.0.0')}"),
        ("清单总项数", "60"),
        ("覆盖维度", "12"),
        ("清单作者", data.get("checklist_author", "缪奇川律师")),
        ("最后更新", data.get("checklist_last_updated", "2026-04-17")),
    ]:
        p = doc.add_paragraph()
        apply_no_indent_paragraph_format(p)
        run = p.add_run(f"{label}:")
        set_run_font(run, size=FONT_SIZE_BODY, bold=True, color=COLOR_BLACK)
        run = p.add_run(value)
        set_run_font(run, size=FONT_SIZE_BODY, color=COLOR_BLACK)

    add_heading_black(doc, "附录 C · 确认层级说明", level=2)
    for label, desc in [
        ("L1 基础级", "AI 识别确定性高,几乎不需要人工复核(如大小写金额不符、错别字)"),
        ("L2 中级", "建议人工复核(如条款前后矛盾、程序性权利陷阱)"),
        ("L3 高级", "必须人工复核(如违约金合理性、数据合规性)"),
    ]:
        p = doc.add_paragraph()
        apply_no_indent_paragraph_format(p)
        run = p.add_run(f"{label}:")
        set_run_font(run, size=FONT_SIZE_BODY, bold=True, color=COLOR_BLACK)
        run = p.add_run(desc)
        set_run_font(run, size=FONT_SIZE_BODY, color=COLOR_BLACK)

    add_heading_black(doc, "附录 D · 规则来源说明", level=2)
    for label, desc in [
        ("标准清单", "由 checklist.md 中 C001-C060 的某一项触发,通用律师共识"),
        ("Playbook 覆盖", "由组织立场手册 playbook 触发,覆盖了 checklist 的默认判断"),
        ("清单外发现", "不属于前两类,由 AI 在审查过程中自主发现的问题"),
    ]:
        p = doc.add_paragraph()
        apply_no_indent_paragraph_format(p)
        run = p.add_run(f"{label}:")
        set_run_font(run, size=FONT_SIZE_BODY, bold=True, color=COLOR_BLACK)
        run = p.add_run(desc)
        set_run_font(run, size=FONT_SIZE_BODY, color=COLOR_BLACK)

    glossary = data.get("glossary", [])
    if glossary:
        add_heading_black(doc, "附录 E · 术语表", level=2)
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["术语 / 简称", "含义"]
        col_widths = [Cm(4), Cm(12)]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.width = col_widths[i]
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = None
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h)
            set_run_font(run, size=FONT_SIZE_SMALL, bold=True, color=COLOR_BLACK)

        for term in glossary:
            row = table.add_row()
            p = row.cells[0].paragraphs[0]
            p.paragraph_format.first_line_indent = None
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(term.get("term", ""))
            set_run_font(run, size=FONT_SIZE_SMALL, bold=True, color=COLOR_BLACK)
            p = row.cells[1].paragraphs[0]
            p.paragraph_format.first_line_indent = None
            run = p.add_run(term.get("definition", ""))
            set_run_font(run, size=FONT_SIZE_SMALL, color=COLOR_BLACK)

        style_data_table(table, has_header=True)
        enforce_cell_formatting(table)


# ============================================================
# 主流程
# ============================================================

def generate_report(data, output_path):
    doc = Document()
    section = doc.sections[0]
    setup_page(section)
    setup_default_paragraph_normal(doc)

    render_cover(doc, data)
    setup_header_footer(section, data)

    render_basic_info(doc, data)
    render_executive_summary(doc, data)
    render_findings_summary(doc, data)
    render_findings_detail(doc, data)
    render_missing_clauses(doc, data)
    render_extra_findings(doc, data)
    render_laws_index(doc, data)
    render_cross_skill_nav(doc, data)
    render_disclaimer(doc, data)
    render_appendix(doc, data)

    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.first_line_indent = None
    run = footer_p.add_run(f"{FOOTER_COPYRIGHT}  |  Skill 版本 v{SKILL_VERSION}")
    set_run_font(run, size=FONT_SIZE_SMALL, color=COLOR_BLACK)

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    print(f"审查报告已保存:{output}")


def main():
    parser = argparse.ArgumentParser(description="生成合同审查报告 Word 文档 v1.0.0")
    parser.add_argument("--output", required=True)
    parser.add_argument("--data", required=True)
    args = parser.parse_args()
    with open(args.data, "r", encoding="utf-8") as f:
        data = json.load(f)
    generate_report(data, args.output)


if __name__ == "__main__":
    main()
