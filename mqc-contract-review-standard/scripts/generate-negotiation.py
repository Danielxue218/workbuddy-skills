#!/usr/bin/env python3
"""
generate-negotiation.py · v1.0.0

生成谈判优先级清单 Word 文档。

核心特性:
1. 严格对齐《参考格式》:16/14/12pt · 1.5行距 · 段前后精确 · 首缩0.85cm · 两端对齐
2. 全正文去色:仅 Tier 的 ■ 方块和 [精准] 前缀保留;其他一切改黑
3. "当前条款 / 建议修改" 堆叠段落 → 三列对照表(含 **...** 红字标注)
4. Plan B 对冲方案文字改黑,关键改动用 **...** 红字
5. 每条 Tier 项新增"规则来源"字段
6. 封面包豪斯风格 · 单外层表格 · 浅蓝 #1387C0 细线
7. 一级/二级标题用 Word 内建 Heading 样式(导航窗格可见)
8. 去掉全文 🎯,改用 [精准] 文字标签
9. 所有表格统一 enforce_cell_formatting(无首缩+垂直居中+段前后对称)

作者:缪奇川
"""

import json
import sys
import argparse
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))

from _common import (
    get_brand_line,
    setup_page, setup_default_paragraph_normal,
    set_run_font, set_cell_shading, set_cell_borders, set_cell_vertical_center,
    setup_table_cell_paragraph, enforce_cell_formatting,
    add_heading_black, add_body_paragraph,
    apply_body_paragraph_format, apply_no_indent_paragraph_format,
    add_comparison_table, _add_revised_text_with_redmark,
    style_header_row_light, apply_data_row_style, style_data_table,
    FONT_CN_BODY, FONT_CN_TITLE, FONT_EN_ALL,
    FONT_SIZE_H1, FONT_SIZE_H2, FONT_SIZE_BODY, FONT_SIZE_SMALL, FONT_SIZE_MINI,
    LINE_SPACING,
    COLOR_RED, COLOR_AMBER, COLOR_GREEN, COLOR_BLACK, COLOR_REVISION_RED,
    HEX_TABLE_HEADER, HEX_TABLE_ZEBRA, HEX_DIVIDER_LINE,
    SYMBOL_RISK_SQUARE,
    TIER_COLOR, TIER_DESC, RULE_SOURCE_LABEL,
    SKILL_VERSION, FOOTER_COPYRIGHT,
)

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================
# 辅助函数
# ============================================================

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def get_tier_color(tier_label):
    if "Tier 1" in tier_label or "tier1" in tier_label.lower():
        return COLOR_RED
    if "Tier 2" in tier_label or "tier2" in tier_label.lower():
        return COLOR_AMBER
    if "Tier 3" in tier_label or "tier3" in tier_label.lower():
        return COLOR_GREEN
    return COLOR_BLACK


def add_divider_line(doc, color_hex=HEX_DIVIDER_LINE):
    """水平细分隔线"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Cm(16)
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        b.set(qn("w:sz"), "0")
        tcBorders.append(b)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), color_hex)
    tcBorders.append(bottom)
    tcPr.append(tcBorders)
    cell.paragraphs[0].text = ""


def setup_header_footer(section, data):
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header_para.add_run(f"{data.get('contract_name', '')}  ·  谈判优先级清单")
    set_run_font(run, size=FONT_SIZE_MINI, color=COLOR_BLACK)

    footer = section.footer
    footer_para = footer.paragraphs[0]
    run = footer_para.add_run(f"审查编号:{data.get('review_number', '')}")
    set_run_font(run, size=FONT_SIZE_MINI, color=COLOR_BLACK)
    footer_para.add_run("\t\t")

    run_pn_label = footer_para.add_run("第 ")
    set_run_font(run_pn_label, size=FONT_SIZE_MINI, color=COLOR_BLACK)
    run_pn = footer_para.add_run()
    _add_field(run_pn, "PAGE")
    set_run_font(run_pn, size=FONT_SIZE_MINI, color=COLOR_BLACK)
    run_sep = footer_para.add_run(" 页 / 共 ")
    set_run_font(run_sep, size=FONT_SIZE_MINI, color=COLOR_BLACK)
    run_tot = footer_para.add_run()
    _add_field(run_tot, "NUMPAGES")
    set_run_font(run_tot, size=FONT_SIZE_MINI, color=COLOR_BLACK)
    run_end = footer_para.add_run(" 页")
    set_run_font(run_end, size=FONT_SIZE_MINI, color=COLOR_BLACK)


def _add_field(run, field_code):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)
    instr = OxmlElement("w:instrText")
    instr.text = field_code
    run._r.append(instr)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


# ============================================================
# 封面·包豪斯风格·浅蓝 #1387C0 · 单外层表格保证单页
# ============================================================

def render_cover(doc, data):
    """封面 · 与审查报告封面同款设计"""
    outer = doc.add_table(rows=14, cols=1)
    outer.alignment = WD_TABLE_ALIGNMENT.LEFT
    outer_cell_width = Cm(15)
    for row in outer.rows:
        row.cells[0].width = outer_cell_width

    row_idx = 0

    _fill_block_cell(outer.cell(row_idx, 0), color_hex="1387C0",
                     width_cm=2.5, height_pt=3)
    row_idx += 1

    _fill_spacer_cell(outer.cell(row_idx, 0), pt=30)
    row_idx += 1

    _fill_text_cell(outer.cell(row_idx, 0), "NEGOTIATION",
                    size=36, bold=True, font_cn="宋体")
    row_idx += 1

    _fill_text_cell(outer.cell(row_idx, 0), "PRIORITY",
                    size=36, bold=True, font_cn="宋体")
    row_idx += 1

    _fill_text_cell(outer.cell(row_idx, 0), "CHECKLIST",
                    size=36, bold=True, font_cn="宋体")
    row_idx += 1

    _fill_spacer_cell(outer.cell(row_idx, 0), pt=8)
    row_idx += 1

    _fill_block_cell(outer.cell(row_idx, 0), color_hex="1387C0",
                     width_cm=2.5, height_pt=2)
    row_idx += 1

    _fill_spacer_cell(outer.cell(row_idx, 0), pt=12)
    row_idx += 1

    _fill_text_cell(outer.cell(row_idx, 0), "谈判优先级清单",
                    size=22, bold=True, font_cn="宋体")
    row_idx += 1

    _fill_block_cell(outer.cell(row_idx, 0), color_hex="1387C0",
                     width_cm=14, height_pt=1)
    row_idx += 1

    _fill_info_table_cell(outer.cell(row_idx, 0), data)
    row_idx += 1

    _fill_block_cell(outer.cell(row_idx, 0), color_hex="1387C0",
                     width_cm=14, height_pt=1)
    row_idx += 1

    _fill_spacer_cell(outer.cell(row_idx, 0), pt=6)
    row_idx += 1

    _fill_text_cell(outer.cell(row_idx, 0),
        get_brand_line(data) or "",
        size=FONT_SIZE_MINI, bold=False, font_cn="宋体")

    _clean_outer_table_cells(outer)
    add_page_break(doc)


def _clean_outer_table_cells(table):
    """外层表格无边框无内边距"""
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._element.get_or_add_tcPr()
            existing = tcPr.find(qn("w:tcBorders"))
            if existing is not None:
                tcPr.remove(existing)
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "right", "bottom"):
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"), "none")
                b.set(qn("w:sz"), "0")
                tcBorders.append(b)
            tcPr.append(tcBorders)
            tcMar = OxmlElement("w:tcMar")
            for side in ("top", "left", "right", "bottom"):
                m = OxmlElement(f"w:{side}")
                m.set(qn("w:w"), "0")
                m.set(qn("w:type"), "dxa")
                tcMar.append(m)
            existing_mar = tcPr.find(qn("w:tcMar"))
            if existing_mar is not None:
                tcPr.remove(existing_mar)
            tcPr.append(tcMar)


def _fill_block_cell(cell, color_hex="1387C0", width_cm=2.5, height_pt=2):
    """用 trH exact 精确控制色块高度"""
    p_existing = cell.paragraphs[0]
    p_existing.paragraph_format.first_line_indent = Cm(0)
    p_existing.paragraph_format.space_before = Pt(0)
    p_existing.paragraph_format.space_after = Pt(0)
    p_existing.paragraph_format.line_spacing = 1.0

    inner_table = cell.add_table(rows=1, cols=1)
    inner_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    ic = inner_table.cell(0, 0)
    ic.width = Cm(width_cm)
    set_cell_shading(ic, color_hex)

    tcPr = ic._element.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "right", "bottom"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        b.set(qn("w:sz"), "0")
        tcBorders.append(b)
    tcPr.append(tcBorders)
    tcMar = OxmlElement("w:tcMar")
    for side in ("top", "left", "right", "bottom"):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), "0")
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)

    p_in = ic.paragraphs[0]
    p_in.paragraph_format.first_line_indent = Cm(0)
    p_in.paragraph_format.space_before = Pt(0)
    p_in.paragraph_format.space_after = Pt(0)
    p_in.paragraph_format.line_spacing = Pt(1)
    p_in.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY

    tr = inner_table.rows[0]._element
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = OxmlElement("w:trPr")
        tr.insert(0, trPr)
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(height_pt * 20)))
    trHeight.set(qn("w:hRule"), "exact")
    trPr.append(trHeight)

    occupant = p_existing.add_run("")
    set_run_font(occupant, size=1, color=COLOR_BLACK)


def _fill_spacer_cell(cell, pt=10):
    p = cell.paragraphs[0]
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(" ")
    set_run_font(run, size=pt, color=COLOR_BLACK)


def _fill_text_cell(cell, text, size=12, bold=False, font_cn="宋体"):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    run = p.add_run(text)
    set_run_font(run, font_name_cn=font_cn, size=size, bold=bold, color=COLOR_BLACK)


def _fill_info_table_cell(cell, data):
    p_existing = cell.paragraphs[0]
    p_existing.paragraph_format.first_line_indent = Cm(0)
    p_existing.paragraph_format.space_before = Pt(0)
    p_existing.paragraph_format.space_after = Pt(0)
    p_existing.paragraph_format.line_spacing = 1.0
    r = p_existing.add_run("")
    set_run_font(r, size=1, color=COLOR_BLACK)

    # 统计 Tier 数量(可选)
    tier1_count = len(data.get("tier1", []))
    tier2_count = len(data.get("tier2", []))
    tier3_count = len(data.get("tier3", []))

    cover_items = [
        ("合同名称", data.get("contract_name", "")),
        ("审查编号", data.get("review_number", "")),
        ("审查立场", data.get("review_stance", "")),
        ("谈判事项", f"Tier 1 · {tier1_count} 项   Tier 2 · {tier2_count} 项   Tier 3 · {tier3_count} 项"),
    ]
    cover_items_f = [(l, v) for l, v in cover_items if v]

    inner = cell.add_table(rows=len(cover_items_f), cols=2)
    inner.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (label, value) in enumerate(cover_items_f):
        c0 = inner.cell(i, 0)
        c1 = inner.cell(i, 1)
        c0.width = Cm(3.2)
        c1.width = Cm(11)
        set_cell_vertical_center(c0)
        set_cell_vertical_center(c1)
        for cc in (c0, c1):
            tcPr = cc._element.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "right", "bottom"):
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"), "none")
                b.set(qn("w:sz"), "0")
                tcBorders.append(b)
            tcPr.append(tcBorders)

        p0 = c0.paragraphs[0]
        p0.paragraph_format.first_line_indent = Cm(0)
        p0.paragraph_format.space_before = Pt(1)
        p0.paragraph_format.space_after = Pt(1)
        p0.paragraph_format.line_spacing = 1.1
        run = p0.add_run(label)
        set_run_font(run, size=FONT_SIZE_BODY, color=COLOR_BLACK)

        p1 = c1.paragraphs[0]
        p1.paragraph_format.first_line_indent = Cm(0)
        p1.paragraph_format.space_before = Pt(1)
        p1.paragraph_format.space_after = Pt(1)
        p1.paragraph_format.line_spacing = 1.1
        run = p1.add_run(value)
        set_run_font(run, size=FONT_SIZE_BODY, bold=True, color=COLOR_BLACK)


# ============================================================
# 一、清单概览
# ============================================================

def render_overview(doc, data):
    add_heading_black(doc, "一、清单概览", level=1)

    tier1_count = len(data.get("tier1", []))
    tier2_count = len(data.get("tier2", []))
    tier3_count = len(data.get("tier3", []))
    total = tier1_count + tier2_count + tier3_count

    add_body_paragraph(doc, f"本合同审查共识别出 {total} 个谈判事项,按优先级分档如下:")

    # 概览表
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["优先级档位", "事项数", "建议处理"]
    col_widths = [Cm(6), Cm(3), Cm(7)]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = col_widths[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, size=FONT_SIZE_BODY, bold=True, color=COLOR_BLACK)

    tiers = [
        ("Tier 1 · 必须回绝", tier1_count, "无法让步", COLOR_RED),
        ("Tier 2 · 建议协商", tier2_count, "重点争取", COLOR_AMBER),
        ("Tier 3 · 可以接受", tier3_count, "作让步筹码", COLOR_GREEN),
    ]
    for tier_label, count, action, color in tiers:
        row = table.add_row()
        cell0 = row.cells[0]
        cell0.width = col_widths[0]
        p = cell0.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(SYMBOL_RISK_SQUARE + " ")
        set_run_font(run, size=FONT_SIZE_BODY, bold=True, color=color)
        run = p.add_run(tier_label)
        set_run_font(run, size=FONT_SIZE_BODY, bold=True, color=COLOR_BLACK)

        cell1 = row.cells[1]
        cell1.width = col_widths[1]
        p = cell1.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(count))
        set_run_font(run, size=FONT_SIZE_BODY, color=COLOR_BLACK)

        cell2 = row.cells[2]
        cell2.width = col_widths[2]
        p = cell2.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(action)
        set_run_font(run, size=FONT_SIZE_BODY, color=COLOR_BLACK)

    style_data_table(table, has_header=True)
    enforce_cell_formatting(table, space_pt=3)

    add_body_paragraph(doc, "使用说明:")
    add_body_paragraph(doc,
        "本清单与同目录下《审查报告》配套使用。每个 Tier 1 / Tier 2 条目都标注了对应的审查报告风险编号(如\"风险 01\")和 checklist 稳定 ID,便于交叉追溯。谈判时应按 Tier 1 → Tier 2 → Tier 3 的顺序推进,Tier 3 作为换取 Tier 2 让步的筹码。")


# ============================================================
# 二-四、Tier 1/2/3 详情(用三列对照表)
# ============================================================

def render_tier_section(doc, data, tier_level, heading_title, tier_intro):
    """渲染一个 Tier 层级的所有条目

    tier_level: 1/2/3
    heading_title: 章节一级标题
    tier_intro: 章节简介
    """
    items = data.get(f"tier{tier_level}", [])
    add_heading_black(doc, heading_title, level=1)

    if tier_intro:
        add_body_paragraph(doc, tier_intro)

    if not items:
        add_body_paragraph(doc, f"本合同审查未发现 Tier {tier_level} 事项。")
        return

    tier_color = [None, COLOR_RED, COLOR_AMBER, COLOR_GREEN][tier_level]

    for idx, item in enumerate(items, 1):
        _render_single_tier_item(doc, tier_level, idx, item, tier_color)


def _render_single_tier_item(doc, tier_level, idx, item, tier_color):
    """渲染一个 Tier 条目"""
    # 条目标题行:"Tier 1-01  ■ 标题"
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(16)
    heading.paragraph_format.space_after = Pt(4)
    heading.paragraph_format.first_line_indent = None

    run = heading.add_run(f"Tier {tier_level}-{idx:02d}   ")
    set_run_font(run, size=FONT_SIZE_H2, bold=True, color=COLOR_BLACK)

    run = heading.add_run(SYMBOL_RISK_SQUARE + "  ")
    set_run_font(run, size=FONT_SIZE_H2, bold=True, color=tier_color)

    run = heading.add_run(item.get("title", ""))
    set_run_font(run, size=FONT_SIZE_H2, bold=True, color=COLOR_BLACK)

    # 元数据行
    rule_source = item.get("rule_source", "checklist")
    source_label = RULE_SOURCE_LABEL.get(rule_source, rule_source)

    meta_items = []
    ref = item.get("risk_report_ref", "")
    if ref and ref != "—":
        meta_items.append(f"对应审查报告:{ref}")
    if item.get("checklist_id"):
        meta_items.append(f"对应清单:{item['checklist_id']}")
    meta_items.append(f"规则来源:{source_label}")

    p = doc.add_paragraph()
    apply_no_indent_paragraph_format(p)
    run = p.add_run("  ·  ".join(meta_items))
    set_run_font(run, size=FONT_SIZE_BODY, color=COLOR_BLACK)

    # 三列对照表
    current_text = item.get("current_text", "")
    suggested_text = item.get("suggested_text", "")

    # 构造"修改理由"——合并 talking_points 为一段
    talking_points = item.get("talking_points", [])
    if isinstance(talking_points, list):
        reason = "  ".join(f"({i+1}) {p}" for i, p in enumerate(talking_points))
    else:
        reason = str(talking_points)

    if current_text and suggested_text:
        p_spacer = doc.add_paragraph()
        p_spacer.paragraph_format.space_after = Pt(0)
        p_spacer.paragraph_format.first_line_indent = None
        add_comparison_table(doc,
            original_text=current_text,
            revised_text=suggested_text,
            reason_text=reason or "(暂无谈判话术)",
            reason_label="谈判话术",
        )

    # 对方可能反驳
    opp = item.get("opponent_rebuttal", None)
    if opp:
        p = doc.add_paragraph()
        apply_no_indent_paragraph_format(p)
        p.paragraph_format.space_before = Pt(6)
        run_pfx = p.add_run("对方可能反驳:")
        set_run_font(run_pfx, size=FONT_SIZE_BODY, bold=True, color=COLOR_BLACK)
        if isinstance(opp, list):
            opp_text = " / ".join(opp)
        else:
            opp_text = opp
        run = p.add_run(opp_text)
        set_run_font(run, size=FONT_SIZE_BODY, color=COLOR_BLACK)

    # 应对要点
    cr = item.get("counter_response", None)
    if cr:
        p = doc.add_paragraph()
        apply_no_indent_paragraph_format(p)
        run_pfx = p.add_run("应对要点:")
        set_run_font(run_pfx, size=FONT_SIZE_BODY, bold=True, color=COLOR_BLACK)
        if isinstance(cr, list):
            cr_text = " / ".join(cr)
        else:
            cr_text = cr
        run = p.add_run(cr_text)
        set_run_font(run, size=FONT_SIZE_BODY, color=COLOR_BLACK)

    # Plan B(Tier 1)或 bottom_line(Tier 2/3)
    plan_b = item.get("plan_b", "")
    bottom_line = item.get("bottom_line", "")

    if plan_b:
        p = doc.add_paragraph()
        apply_no_indent_paragraph_format(p)
        run_pfx = p.add_run("Plan B 对冲方案:")
        set_run_font(run_pfx, size=FONT_SIZE_BODY, bold=True, color=COLOR_BLACK)
        # plan_b 文字支持 **...** 红字标注
        _add_revised_text_with_redmark(p, plan_b)

    if bottom_line:
        p = doc.add_paragraph()
        apply_no_indent_paragraph_format(p)
        run_pfx = p.add_run("底线:")
        set_run_font(run_pfx, size=FONT_SIZE_BODY, bold=True, color=COLOR_BLACK)
        run = p.add_run(bottom_line)
        set_run_font(run, size=FONT_SIZE_BODY, color=COLOR_BLACK)

    # 分隔线
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    add_divider_line(doc)


# ============================================================
# 五、交换矩阵
# ============================================================

def render_exchange_matrix(doc, data):
    matrix = data.get("exchange_matrix", [])
    if not matrix:
        return

    add_heading_black(doc, "五、交换矩阵", level=1)

    add_body_paragraph(doc,
        "谈判策略的核心是\"交换\"——用 Tier 3 的可接受让步,换取对方在 Tier 2 的修改配合。"
        "下表列出本次谈判可操作的关键交换组合。")

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["我方让步", "期待对方修改", "策略说明", "成功率"]
    col_widths = [Cm(3.5), Cm(3.5), Cm(7), Cm(1.8)]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = col_widths[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, size=FONT_SIZE_SMALL, bold=True, color=COLOR_BLACK)

    for item in matrix:
        row = table.add_row()
        vals = [
            item.get("our_concession", ""),
            item.get("their_modification", ""),
            item.get("rationale", ""),
            item.get("success_rate", ""),
        ]
        for i, v in enumerate(vals):
            cell = row.cells[i]
            cell.width = col_widths[i]
            p = cell.paragraphs[0]
            if i == 3:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(v)
            set_run_font(run, size=FONT_SIZE_SMALL, color=COLOR_BLACK)

    style_data_table(table, has_header=True)
    enforce_cell_formatting(table, space_pt=3)


# ============================================================
# 六、使用建议
# ============================================================

def render_usage_advice(doc, data):
    add_heading_black(doc, "六、使用建议", level=1)

    for idx, (title, body) in enumerate([
        ("推进顺序",
         "按 Tier 1 → Tier 2 → Tier 3 的顺序推进谈判。"
         "Tier 1 是不可退让项,必须优先达成;Tier 2 是核心争取项,可动用 Tier 3 作为筹码;"
         "Tier 3 本质是可接受项,在谈判中作为让步筹码使用。"),
        ("时机把握",
         "Tier 1 宜在谈判开局即明确提出,避免后续推诿;"
         "Tier 2 宜在对方已有让步意向后推进;"
         "Tier 3 的让步应该分批次释放,不一次性用尽筹码。"),
        ("Plan B 启用",
         "Plan B 对冲方案是 Tier 1 失败后的保底路径,"
         "启用前应与委托方决策层充分沟通,避免擅自让步。"
         "Plan B 中的关键参数(如金额比例、时限)应作为预设红线,不再降让。"),
        ("结果记录",
         "谈判过程中,应记录每个事项的最终落点:"
         "完全接受 / 部分接受 / 拒绝 / 悬置。这些记录会沉淀为 Playbook 的历史数据,"
         "供后续同类谈判参考。"),
    ], 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.first_line_indent = None
        run = p.add_run(f"{idx}. {title}")
        set_run_font(run, size=FONT_SIZE_BODY, bold=True, color=COLOR_BLACK)
        add_body_paragraph(doc, body)


# ============================================================
# 七、声明
# ============================================================

def render_disclaimer(doc, data):
    add_heading_black(doc, "七、重要声明", level=1)

    declarations = [
        ("1. 辅助性质", "本清单由 AI 工具辅助生成,所有谈判策略仅供参考,必须由执业律师审核后方可使用。本清单不构成法律意见书。"),
        ("2. 场景限制", "本清单基于审查时的合同文本和已知谈判背景。谈判过程中如出现新信息(如对方披露的商业底线、第三方约束),应重新评估本清单的策略。"),
        ("3. 底线保密", "Plan B 对冲方案、底线数据属于我方谈判机密,不应向对方展示本清单。"),
        ("4. 使用范围", "本清单仅供本次谈判委托方及其指定谈判代表使用,未经作者书面同意不得对外披露。"),
    ]

    for title_line, body_text in declarations:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.first_line_indent = None
        run = p.add_run(title_line)
        set_run_font(run, size=FONT_SIZE_BODY, bold=True, color=COLOR_BLACK)
        add_body_paragraph(doc, body_text)


# ============================================================
# 主流程
# ============================================================

def generate_negotiation(data, output_path):
    doc = Document()
    section = doc.sections[0]
    setup_page(section)
    setup_default_paragraph_normal(doc)

    render_cover(doc, data)
    setup_header_footer(section, data)

    render_overview(doc, data)

    render_tier_section(doc, data, 1,
        "二、Tier 1 · 必须回绝(Deal Breaker)",
        "不解决则不能签约。不可作为让步筹码。")

    render_tier_section(doc, data, 2,
        "三、Tier 2 · 建议协商",
        "重点争取修改,但最终可接受折中方案。")

    render_tier_section(doc, data, 3,
        "四、Tier 3 · 可以接受",
        "可以接受原条款。在谈判中作为换取 Tier 2 让步的筹码使用。")

    render_exchange_matrix(doc, data)
    render_usage_advice(doc, data)
    render_disclaimer(doc, data)

    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.first_line_indent = None
    run = footer_p.add_run(f"{FOOTER_COPYRIGHT}  |  Skill 版本 v{SKILL_VERSION}")
    set_run_font(run, size=FONT_SIZE_SMALL, color=COLOR_BLACK)

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    print(f"谈判清单已保存:{output}")


def main():
    parser = argparse.ArgumentParser(description="生成谈判优先级清单 v1.0.0")
    parser.add_argument("--output", required=True)
    parser.add_argument("--data", required=True)
    args = parser.parse_args()
    with open(args.data, "r", encoding="utf-8") as f:
        data = json.load(f)
    generate_negotiation(data, args.output)


if __name__ == "__main__":
    main()
