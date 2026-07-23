#!/usr/bin/env python3
"""
generate-summary.py · v1.0.0

生成合同概要 Word 文档(纯客观信息提取,不含风险分析)。

核心特性:
1. 严格对齐《参考格式》:16/14/12pt · 1.5行距 · 段前后精确 · 首缩0.85cm · 两端对齐
2. 全正文去色:一级/二级标题黑色 · 灰字注释改黑 · 统一黑字呈现
3. 一级/二级标题用 Word 内建 Heading 样式(导航窗格可见)
4. 封面包豪斯风格 · 单外层表格 · 浅蓝 #1387C0 细线
5. 所有表格用 enforce_cell_formatting(无首缩+垂直居中+段前后对称)
6. 合同概要纯客观,无风险符号、无彩色文字

作者:缪奇川
"""

import json
import sys
import argparse
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))

from _common import (
    get_brand_line,
    setup_page, setup_default_paragraph_normal,
    set_run_font, set_cell_shading, set_cell_borders, set_cell_vertical_center,
    setup_table_cell_paragraph, enforce_cell_formatting,
    add_heading_black, add_body_paragraph,
    apply_body_paragraph_format, apply_no_indent_paragraph_format,
    style_header_row_light, apply_data_row_style, style_data_table,
    FONT_CN_BODY, FONT_CN_TITLE, FONT_EN_ALL,
    FONT_SIZE_H1, FONT_SIZE_H2, FONT_SIZE_BODY, FONT_SIZE_SMALL, FONT_SIZE_MINI,
    LINE_SPACING,
    COLOR_BLACK,
    HEX_TABLE_HEADER, HEX_TABLE_ZEBRA, HEX_DIVIDER_LINE,
    SKILL_VERSION, FOOTER_COPYRIGHT,
)

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def setup_header_footer(section, data):
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header_para.add_run(f"{data.get('contract_name', '')}  ·  合同概要")
    set_run_font(run, size=FONT_SIZE_MINI, color=COLOR_BLACK)

    footer = section.footer
    footer_para = footer.paragraphs[0]
    run = footer_para.add_run(f"合同编号:{data.get('contract_number', '')}")
    set_run_font(run, size=FONT_SIZE_MINI, color=COLOR_BLACK)
    footer_para.add_run("\t\t")

    run_pn_label = footer_para.add_run("第 ")
    set_run_font(run_pn_label, size=FONT_SIZE_MINI, color=COLOR_BLACK)
    run_pn = footer_para.add_run()
    _add_field(run_pn, "PAGE")
    set_run_font(run_pn, size=FONT_SIZE_MINI, color=COLOR_BLACK)
    run_sep = footer_para.add_run(" 页 / 共 ")
    set_run_font(run_sep, size=FONT_SIZE_MINI, color=COLOR_BLACK)
    run_tot = footer_para.add_run()
    _add_field(run_tot, "NUMPAGES")
    set_run_font(run_tot, size=FONT_SIZE_MINI, color=COLOR_BLACK)
    run_end = footer_para.add_run(" 页")
    set_run_font(run_end, size=FONT_SIZE_MINI, color=COLOR_BLACK)


def _add_field(run, field_code):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)
    instr = OxmlElement("w:instrText")
    instr.text = field_code
    run._r.append(instr)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


# ============================================================
# 封面
# ============================================================

def render_cover(doc, data):
    outer = doc.add_table(rows=14, cols=1)
    outer.alignment = WD_TABLE_ALIGNMENT.LEFT
    outer_cell_width = Cm(15)
    for row in outer.rows:
        row.cells[0].width = outer_cell_width

    row_idx = 0
    _fill_block_cell(outer.cell(row_idx, 0), color_hex="1387C0",
                     width_cm=2.5, height_pt=3)
    row_idx += 1
    _fill_spacer_cell(outer.cell(row_idx, 0), pt=30)
    row_idx += 1
    _fill_text_cell(outer.cell(row_idx, 0), "CONTRACT", size=36, bold=True, font_cn="宋体")
    row_idx += 1
    _fill_text_cell(outer.cell(row_idx, 0), "SUMMARY", size=36, bold=True, font_cn="宋体")
    row_idx += 1
    _fill_text_cell(outer.cell(row_idx, 0), "BRIEF", size=36, bold=True, font_cn="宋体")
    row_idx += 1
    _fill_spacer_cell(outer.cell(row_idx, 0), pt=8)
    row_idx += 1
    _fill_block_cell(outer.cell(row_idx, 0), color_hex="1387C0",
                     width_cm=2.5, height_pt=2)
    row_idx += 1
    _fill_spacer_cell(outer.cell(row_idx, 0), pt=12)
    row_idx += 1
    _fill_text_cell(outer.cell(row_idx, 0), "合同概要",
                    size=22, bold=True, font_cn="宋体")
    row_idx += 1
    _fill_block_cell(outer.cell(row_idx, 0), color_hex="1387C0",
                     width_cm=14, height_pt=1)
    row_idx += 1
    _fill_info_table_cell(outer.cell(row_idx, 0), data)
    row_idx += 1
    _fill_block_cell(outer.cell(row_idx, 0), color_hex="1387C0",
                     width_cm=14, height_pt=1)
    row_idx += 1
    _fill_spacer_cell(outer.cell(row_idx, 0), pt=6)
    row_idx += 1
    _fill_text_cell(outer.cell(row_idx, 0),
        get_brand_line(data) or "",
        size=FONT_SIZE_MINI, bold=False, font_cn="宋体")

    _clean_outer_table_cells(outer)
    add_page_break(doc)


def _clean_outer_table_cells(table):
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._element.get_or_add_tcPr()
            existing = tcPr.find(qn("w:tcBorders"))
            if existing is not None:
                tcPr.remove(existing)
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "right", "bottom"):
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"), "none")
                b.set(qn("w:sz"), "0")
                tcBorders.append(b)
            tcPr.append(tcBorders)
            tcMar = OxmlElement("w:tcMar")
            for side in ("top", "left", "right", "bottom"):
                m = OxmlElement(f"w:{side}")
                m.set(qn("w:w"), "0")
                m.set(qn("w:type"), "dxa")
                tcMar.append(m)
            existing_mar = tcPr.find(qn("w:tcMar"))
            if existing_mar is not None:
                tcPr.remove(existing_mar)
            tcPr.append(tcMar)


def _fill_block_cell(cell, color_hex="1387C0", width_cm=2.5, height_pt=2):
    p_existing = cell.paragraphs[0]
    p_existing.paragraph_format.first_line_indent = Cm(0)
    p_existing.paragraph_format.space_before = Pt(0)
    p_existing.paragraph_format.space_after = Pt(0)
    p_existing.paragraph_format.line_spacing = 1.0

    inner_table = cell.add_table(rows=1, cols=1)
    inner_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    ic = inner_table.cell(0, 0)
    ic.width = Cm(width_cm)
    set_cell_shading(ic, color_hex)

    tcPr = ic._element.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "right", "bottom"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        b.set(qn("w:sz"), "0")
        tcBorders.append(b)
    tcPr.append(tcBorders)
    tcMar = OxmlElement("w:tcMar")
    for side in ("top", "left", "right", "bottom"):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), "0")
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)

    p_in = ic.paragraphs[0]
    p_in.paragraph_format.first_line_indent = Cm(0)
    p_in.paragraph_format.space_before = Pt(0)
    p_in.paragraph_format.space_after = Pt(0)
    p_in.paragraph_format.line_spacing = Pt(1)
    p_in.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY

    tr = inner_table.rows[0]._element
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = OxmlElement("w:trPr")
        tr.insert(0, trPr)
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(height_pt * 20)))
    trHeight.set(qn("w:hRule"), "exact")
    trPr.append(trHeight)

    occupant = p_existing.add_run("")
    set_run_font(occupant, size=1, color=COLOR_BLACK)


def _fill_spacer_cell(cell, pt=10):
    p = cell.paragraphs[0]
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(" ")
    set_run_font(run, size=pt, color=COLOR_BLACK)


def _fill_text_cell(cell, text, size=12, bold=False, font_cn="宋体"):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    run = p.add_run(text)
    set_run_font(run, font_name_cn=font_cn, size=size, bold=bold, color=COLOR_BLACK)


def _fill_info_table_cell(cell, data):
    p_existing = cell.paragraphs[0]
    p_existing.paragraph_format.first_line_indent = Cm(0)
    p_existing.paragraph_format.space_before = Pt(0)
    p_existing.paragraph_format.space_after = Pt(0)
    p_existing.paragraph_format.line_spacing = 1.0
    r = p_existing.add_run("")
    set_run_font(r, size=1, color=COLOR_BLACK)

    cover_items = [
        ("合同名称", data.get("contract_name", "")),
        ("合同编号", data.get("contract_number", "")),
        ("甲方", data.get("party_a_short", "")),
        ("乙方", data.get("party_b_short", "")),
        ("合同类型", data.get("contract_type", "")),
        ("概要日期", data.get("summary_date", "")),
    ]
    cover_items_f = [(l, v) for l, v in cover_items if v]

    inner = cell.add_table(rows=len(cover_items_f), cols=2)
    inner.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (label, value) in enumerate(cover_items_f):
        c0 = inner.cell(i, 0)
        c1 = inner.cell(i, 1)
        c0.width = Cm(3.2)
        c1.width = Cm(11)
        set_cell_vertical_center(c0)
        set_cell_vertical_center(c1)
        for cc in (c0, c1):
            tcPr = cc._element.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "right", "bottom"):
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"), "none")
                b.set(qn("w:sz"), "0")
                tcBorders.append(b)
            tcPr.append(tcBorders)

        p0 = c0.paragraphs[0]
        p0.paragraph_format.first_line_indent = Cm(0)
        p0.paragraph_format.space_before = Pt(1)
        p0.paragraph_format.space_after = Pt(1)
        p0.paragraph_format.line_spacing = 1.1
        run = p0.add_run(label)
        set_run_font(run, size=FONT_SIZE_BODY, color=COLOR_BLACK)

        p1 = c1.paragraphs[0]
        p1.paragraph_format.first_line_indent = Cm(0)
        p1.paragraph_format.space_before = Pt(1)
        p1.paragraph_format.space_after = Pt(1)
        p1.paragraph_format.line_spacing = 1.1
        run = p1.add_run(value)
        set_run_font(run, size=FONT_SIZE_BODY, bold=True, color=COLOR_BLACK)


# ============================================================
# 通用两列表格渲染(label | value)
# ============================================================

def _render_two_col_table(doc, items, col_widths=None):
    """渲染一个两列(label | value)的表,自动应用 enforce_cell_formatting"""
    if not items:
        return
    if col_widths is None:
        col_widths = [Cm(4), Cm(12)]

    table = doc.add_table(rows=len(items), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (label, value) in enumerate(items):
        c0 = table.cell(i, 0)
        c1 = table.cell(i, 1)
        c0.width = col_widths[0]
        c1.width = col_widths[1]
        set_cell_vertical_center(c0)
        set_cell_vertical_center(c1)
        set_cell_borders(c0)
        set_cell_borders(c1)
        p0 = c0.paragraphs[0]
        run = p0.add_run(label)
        set_run_font(run, size=FONT_SIZE_BODY, bold=True, color=COLOR_BLACK)
        p1 = c1.paragraphs[0]
        run = p1.add_run(value if value else "(合同未约定)")
        set_run_font(run, size=FONT_SIZE_BODY, color=COLOR_BLACK)
    enforce_cell_formatting(table, space_pt=3)


# ============================================================
# 一、合同速览
# ============================================================

def render_overview(doc, data):
    add_heading_black(doc, "一、合同速览", level=1)

    items = [
        ("合同名称", data.get("contract_name", "")),
        ("合同编号", data.get("contract_number", "")),
        ("合同版本", data.get("contract_version", "")),
        ("合同语言", data.get("contract_language", "")),
        ("页数", data.get("page_count", "")),
        ("合同类型", data.get("contract_type", "")),
        ("合同金额", data.get("amount", "")),
        ("合同期限", data.get("term", "")),
        ("签署日期", data.get("sign_date", "")),
        ("生效日期", data.get("effective_date", "")),
        ("争议解决方式", data.get("dispute_method", "")),
    ]
    _render_two_col_table(doc, items, col_widths=[Cm(4), Cm(12)])


# ============================================================
# 二、主体信息
# ============================================================

def render_parties(doc, data):
    add_heading_black(doc, "二、主体信息", level=1)

    # 甲方
    add_heading_black(doc, "2.1 甲方", level=2)
    pa = data.get("party_a_info", {})
    items = [
        ("完整名称", pa.get("full_name", "")),
        ("统一社会信用代码", pa.get("credit_code", "")),
        ("法定代表人", pa.get("legal_representative", "")),
        ("注册地址", pa.get("registered_address", "")),
        ("联系地址", pa.get("contact_address", "")),
        ("项目联系人", pa.get("contact_person", "")),
    ]
    _render_two_col_table(doc, items)

    # 乙方
    add_heading_black(doc, "2.2 乙方", level=2)
    pb = data.get("party_b_info", {})
    items = [
        ("完整名称", pb.get("full_name", "")),
        ("统一社会信用代码", pb.get("credit_code", "")),
        ("法定代表人", pb.get("legal_representative", "")),
        ("注册地址", pb.get("registered_address", "")),
        ("联系地址", pb.get("contact_address", "")),
        ("项目联系人", pb.get("contact_person", "")),
    ]
    _render_two_col_table(doc, items)


# ============================================================
# 三、合同结构索引
# ============================================================

def render_clause_index(doc, data):
    clauses = data.get("clause_index", [])
    if not clauses:
        return

    add_heading_black(doc, "三、合同结构索引", level=1)
    add_body_paragraph(doc, "合同正文的主要条款与附件所在页码索引:")

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["条款编号", "标题", "页码", "备注"]
    col_widths = [Cm(2), Cm(6), Cm(1.5), Cm(6.5)]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = col_widths[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, size=FONT_SIZE_BODY, bold=True, color=COLOR_BLACK)

    for item in clauses:
        row = table.add_row()
        vals = [
            item.get("number", ""),
            item.get("title", ""),
            item.get("page", ""),
            item.get("note", ""),
        ]
        for i, v in enumerate(vals):
            cell = row.cells[i]
            cell.width = col_widths[i]
            p = cell.paragraphs[0]
            if i in (0, 2):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(v)
            set_run_font(run, size=FONT_SIZE_BODY, color=COLOR_BLACK)

    style_data_table(table, has_header=True)
    enforce_cell_formatting(table, space_pt=3)


# ============================================================
# 四、合同标的
# ============================================================

def render_subject(doc, data):
    add_heading_black(doc, "四、合同标的", level=1)
    add_body_paragraph(doc, data.get("subject_description", "(合同未约定)"))


# ============================================================
# 五、权利义务简述
# ============================================================

def render_rights_obligations(doc, data):
    add_heading_black(doc, "五、权利义务简述", level=1)

    pa_ro = data.get("party_a_rights_obligations", {})
    pb_ro = data.get("party_b_rights_obligations", {})

    add_heading_black(doc, "5.1 甲方权利义务", level=2)
    pa_rights = pa_ro.get("rights", [])
    pa_obligations = pa_ro.get("obligations", [])
    items = [
        ("主要权利", " · ".join(pa_rights) if pa_rights else ""),
        ("主要义务", " · ".join(pa_obligations) if pa_obligations else ""),
    ]
    _render_two_col_table(doc, items)

    add_heading_black(doc, "5.2 乙方权利义务", level=2)
    pb_rights = pb_ro.get("rights", [])
    pb_obligations = pb_ro.get("obligations", [])
    items = [
        ("主要权利", " · ".join(pb_rights) if pb_rights else ""),
        ("主要义务", " · ".join(pb_obligations) if pb_obligations else ""),
    ]
    _render_two_col_table(doc, items)


# ============================================================
# 六、付款安排
# ============================================================

def render_payment(doc, data):
    add_heading_black(doc, "六、付款安排", level=1)

    # 6.1 付款阶段
    add_heading_black(doc, "6.1 付款阶段", level=2)
    schedule = data.get("payment_schedule", [])
    if schedule:
        table = doc.add_table(rows=1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["阶段", "金额", "付款条件", "付款期限", "条款"]
        col_widths = [Cm(3), Cm(3.5), Cm(4.5), Cm(3), Cm(2)]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.width = col_widths[i]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h)
            set_run_font(run, size=FONT_SIZE_SMALL, bold=True, color=COLOR_BLACK)

        for item in schedule:
            row = table.add_row()
            vals = [
                item.get("milestone", ""),
                item.get("amount", ""),
                item.get("condition", ""),
                item.get("deadline", ""),
                item.get("clause", ""),
            ]
            for i, v in enumerate(vals):
                cell = row.cells[i]
                cell.width = col_widths[i]
                p = cell.paragraphs[0]
                if i == 4:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(v)
                set_run_font(run, size=FONT_SIZE_SMALL, color=COLOR_BLACK)

        style_data_table(table, has_header=True)
        enforce_cell_formatting(table, space_pt=2)

    # 6.2 收款账户
    add_heading_black(doc, "6.2 收款账户", level=2)
    acc = data.get("payment_account", {})
    items = [
        ("收款方", acc.get("payee", "")),
        ("户名", acc.get("account_name", "")),
        ("开户行", acc.get("bank", "")),
        ("账号", acc.get("account_number", "")),
    ]
    _render_two_col_table(doc, items)

    # 6.3 税务与发票
    add_heading_black(doc, "6.3 税务与发票", level=2)
    tax = data.get("tax_invoice", {})
    items = [
        ("含税情况", tax.get("tax_inclusion", "")),
        ("发票类型", tax.get("invoice_type", "")),
        ("开票期限", tax.get("invoice_deadline", "")),
    ]
    _render_two_col_table(doc, items)


# ============================================================
# 七、关键时间节点
# ============================================================

def render_key_dates(doc, data):
    dates = data.get("key_dates", [])
    if not dates:
        return

    add_heading_black(doc, "七、关键时间节点", level=1)

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["事件", "日期/期限", "责任方", "条款"]
    col_widths = [Cm(4.5), Cm(5), Cm(3.5), Cm(3)]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = col_widths[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, size=FONT_SIZE_BODY, bold=True, color=COLOR_BLACK)

    for item in dates:
        row = table.add_row()
        vals = [
            item.get("event", ""),
            item.get("date", ""),
            item.get("responsible", ""),
            item.get("clause", ""),
        ]
        for i, v in enumerate(vals):
            cell = row.cells[i]
            cell.width = col_widths[i]
            p = cell.paragraphs[0]
            if i in (2, 3):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(v)
            set_run_font(run, size=FONT_SIZE_BODY, color=COLOR_BLACK)

    style_data_table(table, has_header=True)
    enforce_cell_formatting(table, space_pt=3)


# ============================================================
# 八、争议解决
# ============================================================

def render_dispute(doc, data):
    add_heading_black(doc, "八、争议解决", level=1)

    dr = data.get("dispute_resolution", {})
    items = [
        ("解决方式", dr.get("method", "")),
        ("管辖机构", dr.get("jurisdiction", "")),
        ("适用法律", dr.get("governing_law", "")),
    ]
    _render_two_col_table(doc, items)


# ============================================================
# 九、附件清单
# ============================================================

def render_attachments(doc, data):
    attachments = data.get("attachments", [])
    if not attachments:
        return

    add_heading_black(doc, "九、附件清单", level=1)

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["附件名称", "页数", "主要内容", "关联条款"]
    col_widths = [Cm(4), Cm(1.5), Cm(7), Cm(3.5)]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = col_widths[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, size=FONT_SIZE_BODY, bold=True, color=COLOR_BLACK)

    for item in attachments:
        row = table.add_row()
        vals = [
            item.get("name", ""),
            item.get("pages", ""),
            item.get("content", ""),
            item.get("relation", ""),
        ]
        for i, v in enumerate(vals):
            cell = row.cells[i]
            cell.width = col_widths[i]
            p = cell.paragraphs[0]
            if i == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(v)
            set_run_font(run, size=FONT_SIZE_BODY, color=COLOR_BLACK)

    style_data_table(table, has_header=True)
    enforce_cell_formatting(table, space_pt=3)

    # 附件是否已一并签署
    add_body_paragraph(doc, f"附件是否已一并签署:{data.get('attachments_signed', '')}")


# ============================================================
# 十、版本与签署状态
# ============================================================

def render_signing(doc, data):
    signing = data.get("signing_status", {})
    if not signing:
        return

    add_heading_black(doc, "十、版本与签署状态", level=1)

    items = [
        ("甲方用印", signing.get("party_a_seal", "")),
        ("甲方签署日期", signing.get("party_a_date", "")),
        ("乙方用印", signing.get("party_b_seal", "")),
        ("乙方签署日期", signing.get("party_b_date", "")),
        ("骑缝章", signing.get("seal_on_page_edge", "")),
        ("签署平台", signing.get("e_sign_platform", "")),
    ]
    _render_two_col_table(doc, items)


# ============================================================
# 十一、使用说明
# ============================================================

def render_usage(doc, data):
    add_heading_black(doc, "十一、使用说明", level=1)

    for idx, (title, body) in enumerate([
        ("定位",
         "本概要是合同的纯客观信息提取,**不含风险分析和修改建议**。"
         "如需风险分析,请参见同目录下的《审查报告》。"),
        ("未约定字段",
         "合同未明确约定的字段统一标注为\"(合同未约定)\"。不填空不省略。"),
        ("使用场景",
         "本概要适用于:客户评审、内部归档、交接入档、快速了解合同要点等。"),
        ("更新机制",
         "合同如有变更(补充协议、附件变更),应重新生成概要。"
         "本概要版本与合同版本一一对应。"),
    ], 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.first_line_indent = None
        run = p.add_run(f"{idx}. {title}")
        set_run_font(run, size=FONT_SIZE_BODY, bold=True, color=COLOR_BLACK)
        add_body_paragraph(doc, body)


# ============================================================
# 主流程
# ============================================================

def generate_summary(data, output_path):
    doc = Document()
    section = doc.sections[0]
    setup_page(section)
    setup_default_paragraph_normal(doc)

    render_cover(doc, data)
    setup_header_footer(section, data)

    render_overview(doc, data)
    render_parties(doc, data)
    render_clause_index(doc, data)
    render_subject(doc, data)
    render_rights_obligations(doc, data)
    render_payment(doc, data)
    render_key_dates(doc, data)
    render_dispute(doc, data)
    render_attachments(doc, data)
    render_signing(doc, data)
    render_usage(doc, data)

    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.first_line_indent = None
    run = footer_p.add_run(f"{FOOTER_COPYRIGHT}  |  Skill 版本 v{SKILL_VERSION}")
    set_run_font(run, size=FONT_SIZE_SMALL, color=COLOR_BLACK)

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    print(f"合同概要已保存:{output}")


def main():
    parser = argparse.ArgumentParser(description="生成合同概要 v1.0.0")
    parser.add_argument("--output", required=True)
    parser.add_argument("--data", required=True)
    args = parser.parse_args()
    with open(args.data, "r", encoding="utf-8") as f:
        data = json.load(f)
    generate_summary(data, args.output)


if __name__ == "__main__":
    main()
