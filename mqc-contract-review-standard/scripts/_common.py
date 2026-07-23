#!/usr/bin/env python3
"""
_common.py · mqc-contract-review-standard 公共工具模块 · v1.0.0

对齐规范:
- visual-style-guide.md v1.0.0(黑字全局 / 修订红 / 美观三列对照表)
- checklist.md v1.0.0 (稳定 ID / L1-L3 / 立场)
- 《参考格式》(图书正文格式,核心基准)

核心特性:
- 严格对齐《参考格式》:
  一级标题 宋体16pt 黑色 加粗 1.5倍行距 段前16 段后8 首缩0.85cm 两端对齐
  二级标题 宋体14pt 黑色 加粗 1.5倍行距 段前8  段后8 首缩0.85cm 两端对齐
  正文    宋体12pt 黑色      1.5倍行距 段前0  段后0 首缩0.85cm 两端对齐
  英文/数字全部 Times New Roman
- 全正文去色:所有"标题蓝"、"正文灰字"改黑(COLOR_NEUTRAL_* 不再用于正文/标题)
- 新增正红色 COLOR_REVISION_RED = #EE0000 (仅用于修改建议的改动字)
- 新增三列对照表 add_comparison_table(): 美观,垂直居中,两列对照+修改理由整行
- 新增规则来源标签 + 精准审查标记

作者:缪奇川(Miao Qichuan)
"""

# ============================================================
# 一、色板常量
# ============================================================

# ---- 风险等级功能色(仅 ■ 方块使用) ----
COLOR_RED = (255, 0, 0)            # ■ 高风险 / Tier 1
COLOR_AMBER = (255, 153, 0)        # ■ 中风险 / Tier 2
COLOR_GREEN = (0, 166, 80)         # ■ 低风险 / Tier 3

# ---- 基础黑(标题/正文统一) ----
COLOR_BLACK = (0, 0, 0)

# ---- 修订红(仅用于修改建议列内改动字) ----
COLOR_REVISION_RED = (238, 0, 0)   # #EE0000 · 正红,Word 工作习惯

# ---- 扁平风险方块符号 ----
SYMBOL_RISK_SQUARE = "■"
SYMBOL_RISK_CHECK = "✓"
# 注意:精准审查改用文字标签 [精准审查] 而非 emoji 符号,律师文件不要情绪化符号
# SYMBOL_PRECISE 常量保留以便向下兼容,但生成脚本已改用 "[精准审查]" 文字标签
SYMBOL_PRECISE = ""

# ---- 灰色仅用于表格底色 / 分隔线 ----
HEX_TABLE_HEADER = "F1F5F9"
HEX_TABLE_ZEBRA = "F8FAFC"
HEX_TABLE_BORDER = "BFBFBF"
HEX_DIVIDER_LINE = "E0E0E0"

# ============================================================
# 二、文档格式常量(对齐《参考格式》)
# ============================================================

PAGE_MARGIN_TOP = 2.54
PAGE_MARGIN_BOTTOM = 2.54
PAGE_MARGIN_LEFT = 3.17
PAGE_MARGIN_RIGHT = 3.17

FONT_CN_BODY = "宋体"
FONT_CN_TITLE = "宋体"
FONT_EN_ALL = "Times New Roman"

FONT_SIZE_H1 = 16
FONT_SIZE_H2 = 14
FONT_SIZE_BODY = 12
FONT_SIZE_SMALL = 10.5
FONT_SIZE_MINI = 9

LINE_SPACING = 1.5
H1_SPACE_BEFORE = 16
H1_SPACE_AFTER = 8
H2_SPACE_BEFORE = 8
H2_SPACE_AFTER = 8
BODY_SPACE_BEFORE = 0
BODY_SPACE_AFTER = 0
FIRST_LINE_INDENT_CM = 0.85

# ============================================================
# 三、Word 辅助函数
# ============================================================

def _lazy_docx_imports():
    from docx.shared import Pt, Cm, RGBColor
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    return Pt, Cm, RGBColor, qn, WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT


def set_run_font(run, font_name_cn=None, font_name_en=None,
                 size=None, bold=False, color=None):
    """设置文字格式.默认宋体(中)+TNR(英)+黑色+不加粗"""
    Pt, _, RGBColor, qn, *_ = _lazy_docx_imports()
    from docx.oxml import OxmlElement as _Ox

    if font_name_cn is None:
        font_name_cn = FONT_CN_BODY
    if font_name_en is None:
        font_name_en = FONT_EN_ALL

    run.font.name = font_name_en
    if run.element.rPr is None:
        run.element.get_or_add_rPr()
    rPr = run.element.rPr
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = _Ox('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name_cn)
    rFonts.set(qn('w:ascii'), font_name_en)
    rFonts.set(qn('w:hAnsi'), font_name_en)

    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_cell_shading(cell, color_hex):
    _, _, _, qn, *_ = _lazy_docx_imports()
    from docx.oxml import OxmlElement as _Ox
    shading = cell._element.get_or_add_tcPr()
    shading_elem = _Ox('w:shd')
    shading_elem.set(qn('w:val'), 'clear')
    shading_elem.set(qn('w:color'), 'auto')
    shading_elem.set(qn('w:fill'), color_hex)
    shading.append(shading_elem)


def set_cell_vertical_center(cell):
    """设置单元格垂直居中对齐(缪律师本轮要求)"""
    _, _, _, qn, *_ = _lazy_docx_imports()
    from docx.oxml import OxmlElement as _Ox
    tcPr = cell._element.get_or_add_tcPr()
    existing = tcPr.find(qn('w:vAlign'))
    if existing is not None:
        tcPr.remove(existing)
    vAlign = _Ox('w:vAlign')
    vAlign.set(qn('w:val'), 'center')
    tcPr.append(vAlign)


def set_cell_borders(cell, color_hex=HEX_TABLE_BORDER, size_pt=4, sides=None):
    """设置单元格边框.size_pt 为 1/8 pt 单位(4 = 0.5pt)"""
    _, _, _, qn, *_ = _lazy_docx_imports()
    from docx.oxml import OxmlElement as _Ox

    if sides is None:
        sides = ['top', 'left', 'bottom', 'right']

    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = _Ox('w:tcBorders')
        tcPr.append(tcBorders)

    for side in sides:
        existing = tcBorders.find(qn(f'w:{side}'))
        if existing is not None:
            tcBorders.remove(existing)
        border = _Ox(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(size_pt))
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color_hex)
        tcBorders.append(border)


def setup_table_cell_paragraph(para, align="left", space_pt=3):
    """统一表格单元格段落样式:
    - 强制去除首行缩进(避免继承 Normal 样式的 0.85cm)
    - 段前段后对称(避免文字上下不均)
    - 行距 1.15 倍(表格里不要 1.5 倍,否则单元格太高)
    - 两端对齐/居中/左对齐 按需

    参数:
        align: 'left' | 'center' | 'right'
        space_pt: 段前段后对称磅数,默认 3pt
    """
    Pt, Cm, _, _, WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, *_ = _lazy_docx_imports()

    if align == "center":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    pf = para.paragraph_format
    pf.line_spacing = 1.15  # 表格用 1.15 倍(比正文紧凑)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.space_before = Pt(space_pt)
    pf.space_after = Pt(space_pt)
    # 强制去缩进(必须 Cm(0) 而不是 None,才能覆盖 Normal 样式的 0.85cm)
    pf.first_line_indent = Cm(0)


def apply_body_paragraph_format(para):
    """给正文段落应用《参考格式》规范"""
    Pt, Cm, _, _, WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, *_ = _lazy_docx_imports()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = para.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.space_before = Pt(BODY_SPACE_BEFORE)
    pf.space_after = Pt(BODY_SPACE_AFTER)
    pf.first_line_indent = Cm(FIRST_LINE_INDENT_CM)


def apply_no_indent_paragraph_format(para):
    """正文格式但不首缩(元数据行等)"""
    Pt, _, _, _, WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, *_ = _lazy_docx_imports()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = para.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.space_before = Pt(BODY_SPACE_BEFORE)
    pf.space_after = Pt(BODY_SPACE_AFTER)
    pf.first_line_indent = None


def add_heading_black(doc, text, level=1):
    """添加黑色加粗标题,严格《参考格式》规范"""
    Pt, Cm, _, _, WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, *_ = _lazy_docx_imports()

    if level == 1:
        size = FONT_SIZE_H1
        sb = H1_SPACE_BEFORE
        sa = H1_SPACE_AFTER
    else:
        size = FONT_SIZE_H2
        sb = H2_SPACE_BEFORE
        sa = H2_SPACE_AFTER

    # 应用 Word 内建 Heading 样式,这样在"导航窗格"中可见
    # 注意:doc 可以是 Document 或 _Cell,只有 Document 可应用内建样式
    style_name = "Heading 1" if level == 1 else "Heading 2"
    try:
        para = doc.add_paragraph(style=style_name)
    except (KeyError, AttributeError):
        # 如果样式不存在或 doc 不支持,退化为普通段落
        para = doc.add_paragraph()

    run = para.add_run(text)
    # 即使用了 Heading 样式,仍要显式覆盖字体颜色等(否则 Heading 1 默认蓝色)
    set_run_font(run, size=size, bold=True, color=COLOR_BLACK)

    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = para.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.space_before = Pt(sb)
    pf.space_after = Pt(sa)
    pf.first_line_indent = Cm(FIRST_LINE_INDENT_CM)

    return para


def add_body_paragraph(doc, text=None, size=None, bold=False, color=None):
    """添加一段正文(《参考格式》规范)"""
    para = doc.add_paragraph()
    apply_body_paragraph_format(para)
    if text:
        run = para.add_run(text)
        set_run_font(run,
                     size=size or FONT_SIZE_BODY,
                     bold=bold,
                     color=color or COLOR_BLACK)
    return para


def setup_page(section):
    _, Cm, _, _, *_ = _lazy_docx_imports()
    section.top_margin = Cm(PAGE_MARGIN_TOP)
    section.bottom_margin = Cm(PAGE_MARGIN_BOTTOM)
    section.left_margin = Cm(PAGE_MARGIN_LEFT)
    section.right_margin = Cm(PAGE_MARGIN_RIGHT)


def setup_default_paragraph_normal(doc):
    """设置 Normal 样式: 宋体12pt + 1.5行距 + 两端对齐 + 首缩0.85cm"""
    Pt, Cm, _, qn, WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, *_ = _lazy_docx_imports()
    from docx.oxml import OxmlElement as _Ox
    try:
        style = doc.styles['Normal']
        style.font.name = FONT_EN_ALL
        style.font.size = Pt(FONT_SIZE_BODY)
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = _Ox('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), FONT_CN_BODY)
        rFonts.set(qn('w:ascii'), FONT_EN_ALL)
        rFonts.set(qn('w:hAnsi'), FONT_EN_ALL)
        pf = style.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.line_spacing = LINE_SPACING
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.first_line_indent = Cm(FIRST_LINE_INDENT_CM)
    except Exception:
        pass


# ============================================================
# 四、美观三列对照表
# ============================================================

def add_comparison_table(doc, original_text, revised_text, reason_text,
                         reason_label="修改理由"):
    """
    三列对照表:【原文 | 建议】两列 + 【理由】整行
    美观要求:
    - 所有单元格垂直居中
    - 表头浅灰底 + 黑字加粗
    - 边框 #BFBFBF 0.5pt
    - 建议列支持 **...** → 红字(EE0000)

    参数:
        revised_text: 可含 `**改动字**` 标记
        reason_label: 报告/谈判=修改理由,缺失=补充理由
    """
    Pt, Cm, _, qn, WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT = _lazy_docx_imports()

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    col_width = Cm(7.5)
    for col in table.columns:
        for cell in col.cells:
            cell.width = col_width

    # ---- 第1行 表头:条款原文 | 修改建议 ----
    header_row = table.rows[0]
    for i, title in enumerate(["条款原文", "修改建议"]):
        cell = header_row.cells[i]
        cell.width = col_width
        set_cell_shading(cell, HEX_TABLE_HEADER)
        set_cell_vertical_center(cell)
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        setup_table_cell_paragraph(p, align="center", space_pt=4)
        run = p.add_run(title)
        set_run_font(run, size=FONT_SIZE_BODY, bold=True, color=COLOR_BLACK)

    # ---- 第2行 数据:原文 | 建议(含红) ----
    data_row = table.rows[1]

    cell_l = data_row.cells[0]
    cell_l.width = col_width
    set_cell_vertical_center(cell_l)
    set_cell_borders(cell_l)
    p_l = cell_l.paragraphs[0]
    setup_table_cell_paragraph(p_l, align="left", space_pt=4)
    run_l = p_l.add_run(original_text)
    set_run_font(run_l, size=FONT_SIZE_BODY, color=COLOR_BLACK)

    cell_r = data_row.cells[1]
    cell_r.width = col_width
    set_cell_vertical_center(cell_r)
    set_cell_borders(cell_r)
    p_r = cell_r.paragraphs[0]
    setup_table_cell_paragraph(p_r, align="left", space_pt=4)
    _add_revised_text_with_redmark(p_r, revised_text)

    # ---- 第3行 跨列表头:修改理由 ----
    reason_header_row = table.rows[2]
    merged_header = reason_header_row.cells[0].merge(reason_header_row.cells[1])
    set_cell_shading(merged_header, HEX_TABLE_HEADER)
    set_cell_vertical_center(merged_header)
    set_cell_borders(merged_header)
    p_rh = merged_header.paragraphs[0]
    setup_table_cell_paragraph(p_rh, align="center", space_pt=4)
    run_rh = p_rh.add_run(reason_label)
    set_run_font(run_rh, size=FONT_SIZE_BODY, bold=True, color=COLOR_BLACK)

    # ---- 第4行 跨列数据:修改理由内容 ----
    reason_row = table.rows[3]
    merged_reason = reason_row.cells[0].merge(reason_row.cells[1])
    set_cell_vertical_center(merged_reason)
    set_cell_borders(merged_reason)
    p_reason = merged_reason.paragraphs[0]
    setup_table_cell_paragraph(p_reason, align="left", space_pt=4)
    run_reason = p_reason.add_run(reason_text)
    set_run_font(run_reason, size=FONT_SIZE_BODY, color=COLOR_BLACK)

    return table


def _add_revised_text_with_redmark(para, text):
    """解析 **...** 标记: 包裹部分用正红 #EE0000,其余黑色"""
    if not text:
        return

    segments = []
    in_red = False
    buf = ""
    i = 0
    while i < len(text):
        if i + 1 < len(text) and text[i:i+2] == "**":
            if buf:
                segments.append((buf, in_red))
                buf = ""
            in_red = not in_red
            i += 2
        else:
            buf += text[i]
            i += 1
    if buf:
        segments.append((buf, in_red))

    for seg_text, seg_red in segments:
        run = para.add_run(seg_text)
        if seg_red:
            set_run_font(run, size=FONT_SIZE_BODY, color=COLOR_REVISION_RED)
        else:
            set_run_font(run, size=FONT_SIZE_BODY, color=COLOR_BLACK)


# ============================================================
# 五、普通数据表格样式化
# ============================================================

def style_header_row_light(row, column_count):
    """浅灰表头: 浅灰底+黑字加粗+垂直居中+边框"""
    from docx.shared import RGBColor
    for i in range(column_count):
        cell = row.cells[i]
        set_cell_shading(cell, HEX_TABLE_HEADER)
        set_cell_vertical_center(cell)
        set_cell_borders(cell)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(*COLOR_BLACK)


def apply_data_row_style(table, skip_header=True):
    """数据行:垂直居中+边框+斑马纹"""
    start = 1 if skip_header else 0
    for i, row in enumerate(table.rows):
        if i < start:
            continue
        for cell in row.cells:
            set_cell_vertical_center(cell)
            set_cell_borders(cell)
        if (i - start) % 2 == 1:
            for cell in row.cells:
                set_cell_shading(cell, HEX_TABLE_ZEBRA)


def style_data_table(table, has_header=True):
    """一键美化数据表"""
    if not table.rows:
        return
    if has_header:
        style_header_row_light(table.rows[0], len(table.columns))
    apply_data_row_style(table, skip_header=has_header)


def enforce_cell_formatting(table, space_pt=3):
    """一键统一整张表格的单元格格式:
    - 强制所有单元格垂直居中
    - 强制所有段落 first_line_indent = Cm(0)(去除继承的 0.85cm 首缩)
    - 强制所有段落 space_before/after 对称(避免文字上下不居中)
    - 行距 1.15 倍

    对已经设置好对齐方式的段落,保留其原对齐(left/center/right 都兼容)
    只修改缩进、段前后、行距。

    用法:构造完表格、填完内容之后,调用 enforce_cell_formatting(table) 即可。
    """
    Pt, Cm, _, _, _, WD_LINE_SPACING, *_ = _lazy_docx_imports()

    for row in table.rows:
        for cell in row.cells:
            # 垂直居中
            set_cell_vertical_center(cell)
            # 遍历该单元格所有段落
            for para in cell.paragraphs:
                pf = para.paragraph_format
                pf.first_line_indent = Cm(0)
                pf.space_before = Pt(space_pt)
                pf.space_after = Pt(space_pt)
                pf.line_spacing = 1.15
                pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE


# ============================================================
# 六、风险标注映射
# ============================================================

RISK_LEVEL_COLOR = {
    "🔴": COLOR_RED,
    "🟡": COLOR_AMBER,
    "🟢": COLOR_GREEN,
    "■": COLOR_BLACK,
    "高风险": COLOR_RED,
    "中风险": COLOR_AMBER,
    "低风险": COLOR_GREEN,
    "正常": COLOR_BLACK,
}

CONFIRMATION_LEVEL_DESC = {
    "L1": "L1 基础级(AI 可独立确认)",
    "L2": "L2 中级(建议人工复核)",
    "L3": "L3 高级(必须人工复核)",
}

TIER_COLOR = {
    "Tier 1": COLOR_RED,
    "Tier 2": COLOR_AMBER,
    "Tier 3": COLOR_GREEN,
}

TIER_DESC = {
    "Tier 1": "必须回绝 · 不解决则不能签约",
    "Tier 2": "建议协商 · 重点争取修改",
    "Tier 3": "可以接受 · 可作让步筹码",
}

RULE_SOURCE_LABEL = {
    "checklist": "标准清单",
    "playbook": "Playbook 覆盖",
    "extra": "清单外发现",
}


# ============================================================
# 七、品牌标识
# ============================================================

SKILL_NAME = "mqc-contract-review-standard"
SKILL_VERSION = "2.0.0"
AUTHOR = "缪奇川律师"          # 工具作者(Skill 元信息,与经办律师署名无关)
AUTHOR_EN = "Miao Qichuan"
BRAND = "LEGAL AI TOOLMAKER · 法律工具制造者"
FOOTER_COPYRIGHT = f"{BRAND} | {AUTHOR} 出品"
SKILL_SLOGAN = "场景极度垂直 · SOP 极度精简 · 交付极度优雅"

# ------------------------------------------------------------
# 封面底部品牌行(v2.0.0 起可配置)
# 数据 JSON 中提供 "brand_line" 字段即可整体替换;
# 提供空字符串 "" 表示不显示品牌行;不提供则使用默认值。
# ------------------------------------------------------------
DEFAULT_BRAND_LINE = "LEGAL AI TOOLMAKER  ·  法律工具制造者  |  缪奇川律师 出品"


def get_brand_line(data):
    """读取封面品牌行:未配置 → 默认;配置为 "" → None(不显示)。"""
    val = (data or {}).get("brand_line")
    if val is None:
        return DEFAULT_BRAND_LINE
    val = str(val).strip()
    return val or None
